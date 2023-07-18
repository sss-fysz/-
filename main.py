import statistics
import threading
import tkinter as tk
from abc import ABC, abstractmethod
from tkinter import ttk
from tkinter import filedialog
import pickle
import PyPDF2
import mysql
import pandas as pd
import openpyxl
from docx import Document
import numpy as np
from pandas.errors import EmptyDataError
from scipy.stats import norm
import pymysql
from openpyxl import load_workbook
import docx
from openpyxl import Workbook
import matplotlib.pyplot as plt
from sqlalchemy import create_engine
from sqlalchemy.exc import SQLAlchemyError
from pymysql import Connection
from openpyxl.utils.exceptions import InvalidFileException

flows = {}
active_flow_id = None


class Command(ABC):
    @abstractmethod
    def execute(self):
        pass

    @abstractmethod
    def get_id(self):
        pass


class OperateCommand(Command):
    def __init__(self, func, window_id,  *args):
        self.func = func
        self.id = window_id
        self.parameter = args

    def execute(self):
        self.func(*self.parameter)

    def get_id(self):
        return self.id

    def set_parameter(self, *args):
        self.parameter = args


class LogicalCommand(Command):
    def __init__(self, func,window_id, *args):
        self.func = func
        self.window_id = window_id
        self.parameter = args
        self.children_command_list = []
        self.record_children_list = []

    def execute(self):
        for n, command in enumerate(self.children_command_list):
            flows[active_flow_id].text_pad_1.insert(tk.END, f'    <line>{n+1}.')
            command.execute()

    def bind(self, command):
        self.children_command_list.append(command)

    def determine(self):
        return self.func(*self.parameter)

    def get_id(self):
        return self.window_id

    def remove_last_from_logical(self):
        if self.children_command_list:
            self.record_children_list.append(self.children_command_list.pop())


class EndCommand(Command):
    def __init__(self):
        self.window_id = 100

    def execute(self):
        pass

    def get_id(self):
        return self.window_id


class CommandComponent:
    def __init__(self):
        self.command_list = []
        self.is_condition = False
        self.should_stop = False
        self.conditional_command = None
        self.record_list = []

    def bind(self, command):
        if isinstance(command, LogicalCommand):
            self.is_condition = True
            self.conditional_command = command
            self.command_list.append(command)
        elif self.is_condition and isinstance(command, EndCommand):
            self.is_condition = False
            self.conditional_command = None
            self.command_list.append(command)
        elif self.is_condition and self.conditional_command:
            self.conditional_command.bind(command)
        else:
            self.command_list.append(command)

    def click(self):
        for n, command in enumerate(self.command_list):
            if self.should_stop:
                break
            elif isinstance(command, LogicalCommand):
                if command.determine():
                    flows[active_flow_id].text_pad_1.insert(tk.END, f'<Line{n+1}>:逻辑通过，子函数开始运行\n')
                    command.execute()
                else:
                    flows[active_flow_id].text_pad_1.insert(tk.END, f'<Line{n+1}>:逻辑不通过\n','red')
                    flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
            elif isinstance(command, EndCommand):
                pass
            else:
                flows[active_flow_id].text_pad_1.insert(tk.END, f'<Line{n+1}>:')
                command.execute()

    def remove_last(self):
        if self.command_list:
            last_command = self.command_list[-1]
            if isinstance(last_command, LogicalCommand):
                if len(last_command.children_command_list) > 0:
                    last_command.remove_last_from_logical()
                else:
                    self.record_list.append(last_command)
                    self.command_list.pop()
            else:
                self.record_list.append(last_command)
                self.command_list.pop()

    def clear(self):
        self.command_list.clear()
        self.record_list.clear()

    def save_flow(self, file_path):
        command_data = []
        for command in self.command_list:
            if isinstance(command, LogicalCommand):
                command_data.append((command, command.children_command_list))
            else:
                command_data.append(command)

        with open(file_path, 'wb') as f:
            pickle.dump(command_data, f)
        print(f"Saved flow to {file_path}")

    def load_flow(self, file_path):
        with open(file_path, 'rb') as f:
            command_data = pickle.load(f)

        self.command_list = []
        for data in command_data:
            if isinstance(data, tuple) and isinstance(data[0], LogicalCommand):
                command, children_command_list = data
                command.children_command_list = children_command_list
            else:
                command = data
            self.command_list.append(command)


class Flow:
    def __init__(self):
        self.commandcomponent = CommandComponent()
        self.text_pad = None
        self.text_pad_1 = None
        self.logical_status = False

    def active(self):
        self.text_pad = tk.Text(info_frame, width=58, height=25, autoseparators=False, undo=True, maxundo=100)
        self.text_pad.pack(fill=tk.BOTH, expand=True)
        self.text_pad_1 = tk.Text(sendMessage_frame_Children, width=58, height=25, autoseparators=False, undo=True, maxundo=100)
        self.text_pad_1.pack(fill=tk.BOTH, expand=True)
        
    def hide(self):
        self.text_pad.pack_forget()
        self.text_pad_1.pack_forget()
        
        
def create_flow():
    global active_flow_id

    flow_id = len(flows) + 1
    flows[flow_id] = Flow()
    flows[flow_id].active()

    if active_flow_id is not None:
        flows[active_flow_id].hide()

    active_flow_id = flow_id


def switch_flow():
    global active_flow_id

    if active_flow_id is not None:
        flows[active_flow_id].hide()
    if active_flow_id is None:
        pass
    else:
        active_flow_id = (active_flow_id % len(flows)) + 1
        flows[active_flow_id].text_pad.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        flows[active_flow_id].text_pad_1.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)


def replace_excel_values(filepath ,sheet_name,start_cell,end_cell,old_value,new_value):
    """
         功能：将指定 Excel 文件中指定表格、指定位置范围内的某个特定数据进行替换。
         对应窗口：ExcelFunction3_window
         参数列表:  :param self.replace_excel_values_1start_cell: 起点坐标（字符串类型），例如 'A1'。
                  :param self.replace_excel_values_1end_cell: 终点坐标（字符串类型），例如 'H10'。
                  :param self.replace_excel_values_1old_value: 某个特定数据的值。
                  :param self.replace_excel_values_1new_value: 需要替换成的新值。
         报错提示：由于Excel的不加双引号的都是数值类型
                1.范围不存在
         :return:
    """
    # 用于检测是不是字符串
    def is_numeric_string(s):
        # 尝试将字符串转换为浮点数类型，如果出现异常则说明不是数字型字符串
        try:
            float(s)
            return True
        except ValueError:
            return False

    try:
        # 读取 Excel 文件到 Pandas数据框中
        df = pd.read_excel(filepath, sheet_name=sheet_name, header=None)
        # 解析起始和结束单元格的行列索引
        start_row, start_col = int(start_cell[1:]) - 1, ord(
            start_cell[0]) - 65
        end_row, end_col = int(end_cell[1:]) - 1, ord(
            end_cell[0]) - 65
        subset = df.loc[start_row:end_row, start_col:end_col]
        # 判断是否是数字字符串
        if isinstance(old_value, str):
            if old_value.isdigit() or is_numeric_string(old_value):
                num1 = float(old_value)
                if num1 == int(num1):  # 判断其是浮点数还是整数
                    num1 = int(num1)
                old_value = num1
            else:
                pass

        if isinstance(new_value, str):
            if new_value.isdigit() or is_numeric_string(new_value):
                num2 = float(new_value)
                if num2 == int(num2):
                    num2 = int(num2)
                new_value = num2
            else:
                pass
        # 利用 replace 方法进行替换操作
        subset.replace(old_value, new_value, inplace=True)
        # 将原数据框对应部分赋值为修改后的结果
        df.loc[start_row:end_row, start_col:end_col] = subset
        # 将修改后的数据框写回 Excel 文件中
        df.to_excel(filepath, sheet_name=sheet_name, index=False, header=None)
        flows[active_flow_id].text_pad_1.insert(tk.END, '指定范围替换成功！\n')
    except FileNotFoundError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"指定范围替换失败:未找到对应文件\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, '指定范围替换失败:选取范围超出表格范围\n', "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")


def insert_excel_to_word(filepath, sheet_name, word_filepath, start_coordinate=-1, end_coordinate=-1):
    """
        将指定路径的指定范围内的excel数据插入指定路径的word文件
    """
    try:
        # 打开 Excel 文件并获取工作表
        work_book = openpyxl.load_workbook(filepath,sheet_name)
        work_sheet = work_book.active
        if start_coordinate == '-1' and end_coordinate == '-1':
            workbook = load_workbook(filepath)
            start_coordinate = workbook.active.cell(row=workbook.active.min_row,
                                                    column=workbook.active.min_column).coordinate
            end_coordinate = workbook.active.cell(row=workbook.active.max_row,
                                                  column=workbook.active.max_column).coordinate
        # 获取表格数据
        data = [
            [cell.value for cell in row]
            for row in work_sheet[start_coordinate:end_coordinate]
        ]

        # 获取行和列的数量
        num_rows = len(data)
        num_cols = len(data[0])

        # 建立新的 Word 文档并添加表格
        document = Document(word_filepath)
        table = document.add_table(rows=num_rows + 1, cols=num_cols)
        table.style = 'Table Grid'
        # 向表格中添加列数据（省略第一行的列名）
        for i, row_data in enumerate(data):
            for j, cell_data in enumerate(row_data):
                table.cell(i + 1, j).text = str(cell_data)

        # 保存 Word 文件到磁盘
        document.save(word_filepath)
        flows[active_flow_id].text_pad_1.insert(tk.END, '指定范围excel数据插入成功！\n')
    except openpyxl.utils.exceptions.InvalidFileException as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"指定范围excel数据插入失败:未找到对应文件\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"指定范围excel数据插入失败:{e}\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")


def generate_bar_chart(filepath ,sheet_name,top_left_coord,bottom_right_coord,picture_path):
    """
        功能描述：
        将指定范围内的excel表格内的数据，以左上角坐标所在行为数值列，以所在列为行索引(都不包括左上角坐标本身)
        生成一个路径为文件路径+"-picture"的jpg格式的文件。
    """
    try:
        # 读取Excel文件
        df = pd.read_excel(filepath, sheet_name=sheet_name, header=None)
        # 默认为全局
        if top_left_coord == "-1" and bottom_right_coord == "-1":
            # 获得表格的行数和列数
            num_rows, num_cols = df.shape
            # 左上角坐标为 (0, 0)，右下角坐标为 (num_rows-1, num_cols-1)
            top_left_coord = 'A1'
            bottom_right_coord = chr(num_cols + 64) + str(num_rows)
        # 将坐标转换为数字
        top_left_col = ord(top_left_coord[0]) - 65
        top_left_row = int(top_left_coord[1:]) - 1
        bottom_right_col = ord(bottom_right_coord[0]) - 65
        bottom_right_row = int(bottom_right_coord[1:]) - 1
        # 选取指定范围
        df = df.iloc[top_left_row:bottom_right_row + 1, top_left_col:bottom_right_col + 1]
        # 设置行索引和列名
        df.columns = df.iloc[0, :]
        df = df.iloc[1:, :]
        df.index = df.iloc[:, 0]
        df.drop(df.columns[0], axis=1, inplace=True)
        df.index.name, df.columns.name = None, None
        # 绘制柱状图
        df.plot(kind='bar', rot=0)
        if picture_path == '':
            plt.savefig(filepath.replace(".xlsx", "-picture.jpg"))
            flows[active_flow_id].text_pad_1.insert(tk.END, '根据表格数据生成多柱状图成功！\n')
        else:
            plt.savefig(picture_path)
            flows[active_flow_id].text_pad_1.insert(tk.END, '根据表格数据生成多柱状图成功！\n')
    except FileNotFoundError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, '根据表格数据生成多柱状图失败:没有找到对应文件\n', "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f'ERROR: 根据表格数据生成多柱状图失败{e}\n', "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")



def insert_word_to_excel(word_path, excel_filepath):
    """
        将指定路径的word文件的表格插入指定路径的excel文件
    """
    try:
        # 打开 Word 文档
        doc = docx.Document(word_path)
        # 获取所有表格的内容
        tables = [table for table in doc.tables]
        # 创建一个新的 Excel 工作簿
        wb = Workbook()
        ws = wb.active
        # 循环遍历每个表格，并将其转换为列表形式
        for table in tables:
            table_data = []
            for i, row in enumerate(table.rows):
                # 注意：由于第一行表头可能包含空单元格，
                # 在循环第一行时需要手动计算列数。
                if i == 0:
                    col_count = len(row.cells)
                row_data = []
                for cell in row.cells:
                    value = cell.text.strip()
                    row_data.append(value)
                table_data.append(row_data)
            # 表格与表格之间插入分隔符行
            if len(tables) > 1:
                separator = ['***'] * col_count
                table_data.append(separator)
            # 将表格数据逐行写入工作表中
            for row_data in table_data:
                ws.append(row_data)
        # 保存工作簿到文件中
        wb.save(excel_filepath)
        flows[active_flow_id].text_pad_1.insert(tk.END, '指定word表格插入excel成功！\n')
    except FileNotFoundError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, ' 将word文件的表格插入excel文件失败:没有找到对应文件\n', "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, ' 将word文件的表格插入excel文件失败:{e}\n', "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")


def merge_Excel_files(Summary_document_path,*file_paths):

    try:
        if file_paths:
            df_list = []
            for file_path in file_paths:
                df_list.append(pd.read_excel(file_path))
            # 将多个 Dataframe 进行合并
            combined_df = pd.concat(df_list, ignore_index=True)
            # 将合并后的结果转换成 JSON 格式
            json_data = combined_df.to_json(orient='records')
            df = pd.read_json(json_data)
            # 将 DataFrame 转换成 Excel 文件并保存
            writer = pd.ExcelWriter(Summary_document_path)
            df.to_excel(writer, index=False,header=None)
            writer._save()
            flows[active_flow_id].text_pad_1.insert(tk.END, '合并文件成功\n')

    except FileNotFoundError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, ' 合并文件失败:没有找到对应文件\n', "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, '合并文件失败：{e}\n', "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")



def fill_missing_data_in_excel(filepath, sheetname, left_top, right_bottom):
    print(filepath, sheetname, left_top, right_bottom)
    print(type(filepath), type(sheetname), type(left_top), type(right_bottom))
    try:
        def convert_coordinate(coordinate):
            """
            将坐标字符串转换为数字索引
            coordinate: 坐标字符串，如 "A1"
            返回值: 数字索引的元组，如 (0, 0)
            """
            # 将字母部分转换为列索引
            column = ord(coordinate[0].upper()) - ord('A')
            # 将数字部分转换为行索引
            row = int(coordinate[1:]) - 1

            return row, column

        def em_algorithm(data):
            """
            对于一列数据，采用最大似然方法求解其中的缺失部分，并对其进行插补
            data: 数据序列，包含缺失值和非缺失值。
                  如果序列中存在缺失值，则需要用 None 来代替缺失的数值。
            return : 插补后的数据序列，缺失值填充为 EM 算法插补后的值（均值）。
            """

            # 将数据转换为 Numpy 数组
            x = np.array(data, dtype=np.float64, copy=True)
            # 检查数据中是否存在无效值（如负数、无穷大等）
            x = np.where(np.isfinite(x), x, np.nan)

            # 保留原始数据长度
            N = len(x)

            # 构造相关变量
            # 检查数据中的缺失值
            missing_values = np.isnan(x)

            # 如果存在缺失值，则将其填充为均值
            if np.any(missing_values):
                if np.all(missing_values):
                    return x  # 数据全部缺失，无法进行插补
                x[missing_values] = np.nanmean(x[~missing_values])
            # mean : 均值
            # std : 标准差
            # E_xt_yt : 确定值 X(t) 对应的 y(t) 的条件概率分布的期望, 即是 y(t) 关于当前的 X(t) 的条件期望值.
            mean = np.nanmean(x)
            std = np.nanstd(x)
            E_xt_yt = np.empty(N)

            # 初始化参数
            alpha = np.ones(2)
            alpha /= np.sum(alpha)
            log_P = np.full((N, 2), -np.Inf)

            # 循环迭代直至收敛
            while True:

                # 使用均值和标准差估算正态分布函数的参数
                log_likelihood = np.nan_to_num(norm.logpdf(x[:, None], mean, std + 1e-6))

                # 计算 E_step
                for i in range(N):
                    if np.isnan(x[i]):
                        log_P[i] = np.log(alpha)
                    else:
                        log_P[i] = log_likelihood[i]
                log_P -= np.max(log_P, axis=1)[:, None]
                P = np.exp(log_P)
                P /= np.sum(P, axis=1)[:, None]

                # 计算 M_step
                alpha = np.sum(P, axis=0)
                alpha /= np.sum(alpha)

                # 统计期望数
                E_xt_yt.fill(0)
                for i in range(N):
                    E_xt_yt[i] = alpha[1] * (mean / (std ** 2 + 1e-6)) * \
                                 norm.pdf((x[i] - mean) / (std + 1e-6)) if not np.isnan(x[i]) else alpha[1] * mean
                # 判断迭代是否收敛
                new_mean = np.sum(E_xt_yt)
                if np.abs(new_mean - mean) < 1e-6:
                    break
                mean = new_mean

            # 填充缺失值为 EM 算法插补后的数据
            for i in range(N):
                if np.isnan(x[i]):
                    x[i] = E_xt_yt[i]

            return x
        """
        通过最大期望算法对Excel表中的空白内容进行插补，在指定的范围内执行插补操作。
        filepath:   Excel文件路径
        sheetname:  表名
        left_top:   矩形区域左上角坐标，如 "A1"
        right_bottom: 矩形区域右下角坐标，如 "B1"
        """
        # 将坐标转换为数字索引
        left_top_index = convert_coordinate(left_top)
        right_bottom_index = convert_coordinate(right_bottom)

        # 加载 Excel 文件
        df = pd.read_excel(filepath, sheet_name=sheetname, na_values='', header=None)

        # 获取需要处理的区域
        data_df = df.iloc[left_top_index[0]:right_bottom_index[0], left_top_index[1]:right_bottom_index[1]]

        # 对空白内容的缺失值进行插补
        for col in data_df.columns:
            column_data = data_df[col]

            # 处理数据中的异常值
            column_data = column_data.apply(lambda x: np.nan if x == '' or not np.isfinite(x) else x)

            em_data = em_algorithm(column_data.values)
            data_df[col] = pd.DataFrame(em_data)

        # 将新的数据写回 Excel 文件
        df.iloc[left_top_index[0]:right_bottom_index[0], left_top_index[1]:right_bottom_index[1]] = data_df
        df.to_excel(filepath, sheet_name=sheetname, index=False, header=None)
        flows[active_flow_id].text_pad_1.insert(tk.END, '最大似然法填充成功\n')
    except IndexError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f' 最大似然法填充失败:没有找到对应文件或选取范围超过表格范围\n{e}', "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f'最大似然法填充失败\n:{e}', "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")


def process_excel_data(file_path, sheet_name, column_name1, weight, column_name2, operation):
    def column_name_to_index(column_name):
        """
        将列号转换为列索引。

        参数：
        column_name：字符串类型，列号，如"A"、"B"、"C"。

        返回值：
        整数类型，列索引。
        """
        column_index = ord(column_name.upper()) - ord('A')
        return column_index

    try:
        df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)

        # 将列号转换为列索引
        column_index1 = column_name_to_index(column_name1)
        column_index2 = column_name_to_index(column_name2)

        max_rows = max(df.iloc[:, column_index1].shape[0], df.iloc[:, column_index2].shape[0])

        if column_index2 >= df.shape[1]:
            df.insert(loc=column_index2, column='', value=0)
            df.iloc[:max_rows, column_index2] = df.iloc[:max_rows, column_index1] * float(weight) + df.iloc[:max_rows,
                                                                                                    column_index2]
        else:
            df.iloc[:max_rows, column_index2] = df.iloc[:max_rows, column_index1] * float(weight) + df.iloc[:max_rows,
                                                                                                    column_index2]

        df.to_excel(file_path, sheet_name=sheet_name, index=False, header=None)
        flows[active_flow_id].text_pad_1.insert(tk.END, '运算填充成功\n')
    except FileNotFoundError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, ' 运算填充失败:没有找到对应文件或选取范围超过表格范围\n', "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f'运算填充失败:{e}\n', "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")


def search(excel_filepath, sheet, search_value,my_type):
    print(excel_filepath, sheet, search_value,my_type)
    print(type(excel_filepath), type(sheet), type(search_value),type(my_type))
    try:
        df = pd.read_excel(excel_filepath, sheet_name=sheet)
        # 检查是否存在数值类型的 search_value
        if my_type=="int":
            search_value = int(search_value)
        elif my_type == "float":
            search_value = float(search_value)
        elif my_type == "string":
            pass
        print(type(search_value))
        if df.isin([search_value]).any().any():
            return True
        # 检查是否存在字符串类型的 search_value
        elif df.isin([str(search_value)]).any().any():
            return True
        else:
            return False
    except Exception as e:
        print(e)

def export_data_to_excel(host, port, user, password, database, table, output_file):

    try:
        engine = create_engine(f"mysql+pymysql://{user}:{password}@{host}:{int(port)}/{database}")
        # Use pandas to read data from MySQL
        df = pd.read_sql_table(table, con=engine)

        # Export data to Excel file
        with pd.ExcelWriter(output_file, mode='a', engine="openpyxl") as writer:
            if table in writer.book.sheetnames:
                writer.book.remove(writer.book[table])
            df.to_excel(writer, index=False, sheet_name=table)
            writer._save()

        flows[active_flow_id].text_pad_1.insert(tk.END, f':从mysql读取数据到excel成功\n')
    except ValueError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f" 从mysql读取数据到excel失败：参数错误\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except pymysql.OperationalError as oe:
        flows[active_flow_id].text_pad_1.insert(tk.END,f" 从mysql读取数据到excel失败：mysql操作错误\n","red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except pymysql.ProgrammingError as pe:
        flows[active_flow_id].text_pad_1.insert(tk.END,f"从mysql读取数据到excel失败：mysql项目错误\n","red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except pymysql.DatabaseError as de:
        flows[active_flow_id].text_pad_1.insert(tk.END,f"从mysql读取数据到excel失败：数据库错误\n","red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except EmptyDataError as ede:
        flows[active_flow_id].text_pad_1.insert(tk.END,"从mysql读取数据到excel失败: No data found in the specified table.\n","red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except SQLAlchemyError as se:
        flows[active_flow_id].text_pad_1.insert(tk.END,f"从mysql读取数据到excel失败：数据库无法链接，检查一下传入参数\n","red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END,f"从mysql读取数据到excel失败：其他错误 \n ","red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")


def export_data_to_word(host, port, user, password, database, table, output_file):
    try:
        # 连接到MySQL数据库
        connection = Connection(
            host=host,
            port=int(port),
            user=user,
            password=password,
            database=database
        )
    except ValueError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f" 从mysql读取数据到word导出失败:数据库链接错误检查一下参数\n ", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
        return
    try:
        # 查询表数据
        cursor = connection.cursor()
        cursor.execute(f"SELECT * FROM {table}")
        rows = cursor.fetchall()
        column_names = [desc[0] for desc in cursor.description]
    except mysql.connector.Error as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f" 从mysql读取数据到word导出失败:数据表查询失败\n ", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
        return
    finally:
        cursor.close()

    try:
        # 创建Word文档
        doc = Document()

        # 添加表格
        table = doc.add_table(rows=1, cols=len(column_names), style='Table Grid')

        # 添加表头
        for i, column_name in enumerate(column_names):
            table.cell(0, i).text = column_name

        # 添加数据
        for row in rows:
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

        # 保存Word文档
        doc.save(output_file)
        flows[active_flow_id].text_pad_1.insert(tk.END, f" 从mysql读取数据到word成功导出 \n ", "red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f" 从mysql读取数据到word导出失败:数据导出失败 \n ", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    finally:
        connection.close()


def pdf_convert_word(pdf_filepath, start_page, end_page,output_file):
    # 创建一个新的Word文档对象
    try:
        doc = Document()

        # 打开PDF文件并读取指定页码范围
        with open(pdf_filepath, 'rb') as pdf:
            reader = PyPDF2.PdfReader(pdf)
            # 获取第 i 页的内容
            # text = "\n".join([reader.pages[i].extract_text() for i in range(start_page - 1, end_page)])
            # 将提取的内容作为新段落添加到Word文档中
            #doc.add_paragraph(text)
            for i in range(start_page-1, end_page):
                page = reader.pages[i]
                text = page.extract_text()
                doc.add_paragraph(f"第{i}页pdf转换结果\n")
                # 将提取的内容作为新段落添加到Word文档中
                doc.add_paragraph(text)
        # 保存Word文档
        doc.save(output_file)
        flows[active_flow_id].text_pad_1.insert(tk.END,"从pdf中读取数据到excel中成功\n")
    except FileNotFoundError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END,"从pdf中读取数据到excel中失败：指定的PDF文件路径不存在或者无法被访问\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except PyPDF2.utils.PdfReadError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END,"从pdf中读取数据到excel中失败：读取PDF文件时出现了错误，例如PDF文件被加密或损坏。\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except TypeError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END,"从pdf中读取数据到excel中失败：在处理PDF文件页码时可能会出错。\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")


def merge_excel_sheets(excel_file1_path, excel_file2_path, sheet_name1, sheet_name2, merged_file_path, merged_sheet_name):
    try:
        # 加载两个Excel文件
        wb1 = openpyxl.load_workbook(excel_file1_path)
        wb2 = openpyxl.load_workbook(excel_file2_path)

        # 获取两个表格
        sheet1 = wb1[sheet_name1]
        sheet2 = wb2[sheet_name2]

        # 创建一个新的Excel文件用于存储合并后的数据
        merged_wb = openpyxl.Workbook()
        merged_sheet = merged_wb.active
        merged_sheet.title = merged_sheet_name

        # 遍历sheet1和sheet2的数据并将其添加到合并后的表格中
        for sheet in [sheet1, sheet2]:
            for row in sheet.iter_rows():
                row_data = [cell.value for cell in row]
                merged_sheet.append(row_data)

        # 保存合并后的Excel文件
        merged_wb.save(merged_file_path)
        flows[active_flow_id].text_pad_1.insert(tk.END, "Excel文件表格合并成功\n")
    except FileNotFoundError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, "Excel文件表格合并失败:文件没找到\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except InvalidFileException as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, "Excel文件表格合并失败:无效的文件格式\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except KeyError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, "Excel文件表格合并失败:表格名称错误\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, "Excel文件表格合并失败:未知错误\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")


def merge_word_files(file1_path, file1_start, file1_end, file2_path, file2_start, file2_end, output_path):
    try:
        file1_start, file1_end = int(file1_start), int(file1_end)
        file2_start, file2_end = int(file2_start), int(file2_end)
        # 打开文件1和文件2
        doc1 = Document(file1_path)
        doc2 = Document(file2_path)

        # 创建一个新的Word文档用于存储合并及其后的内容
        merged_doc = Document()

        # 将文件1的指定范围内的内容添加到合并后的文件中
        for i in range(file1_start - 1, file1_end):
            paragraph = doc1.paragraphs[i]
            new_paragraph = merged_doc.add_paragraph(paragraph.text, paragraph.style)
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline

        # 将文件2的指定范围内的内容添加到合并后的文件中
        for i in range(file2_start - 1, file2_end):
            paragraph = doc2.paragraphs[i]
            new_paragraph = merged_doc.add_paragraph(paragraph.text, paragraph.style)
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline

        # 保存合并后的文件
        merged_doc.save(output_path)
        flows[active_flow_id].text_pad_1.insert(tk.END, "word文件合并成功\n")
    except FileNotFoundError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, "word文件合并失败:文件没找到\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"word文件合并失败:其他原因{e}\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")


def merge_pdf_files(pdf1_path, pdf1_start, pdf1_end, pdf2_path, pdf2_start, pdf2_end, output_path):
    try:
        # 创建一个PdfFileReader对象，用于读取pdf文件1和pdf文件2
        pdf1_reader = PyPDF2.PdfFileReader(pdf1_path)
        pdf2_reader = PyPDF2.PdfFileReader(pdf2_path)

        # 创建一个PdfFileWriter对象，用于写入合并后的pdf文件
        pdf_writer = PyPDF2.PdfFileWriter()

        # 合并pdf文件1的指定范围
        for page_num in range(pdf1_start - 1, pdf1_end):
            page = pdf1_reader.getPage(page_num)
            pdf_writer.addPage(page)

        # 合并pdf文件2的指定范围
        for page_num in range(pdf2_start - 1, pdf2_end):
            page = pdf2_reader.getPage(page_num)
            pdf_writer.addPage(page)

        # 将合并后的pdf文件写入到指定的输出文件
        with open(output_path, 'wb') as output_file:
            pdf_writer.write(output_file)
        flows[active_flow_id].text_pad_1.insert(tk.END, "pdf文件合并并成功")
    except FileNotFoundError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, "pdf文件合并失败: 当指定的输入文件之一不存在时引发的错误\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except OSError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, "pdf文件合并失败: 权限问题或无法打开文件\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END,
                          f"pdf文件合并失败: 其他意外错误{e}\n",
                          "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")



def EM_data_in_excel(filepath, sheetname, left_top, right_bottom):
    try:
        def convert_coordinate(coordinate):
            """
            将坐标字符串转换为数字索引
            coordinate: 坐标字符串，如 "A1"
            返回值: 数字索引的元组，如 (0, 0)
            """
            # 将字母部分转换为列索引
            column = ord(coordinate[0].upper()) - ord('A')
            # 将数字部分转换为行索引
            row = int(coordinate[1:]) - 1

            return row, column

        def em_algorithm(data):
            """
            使用最大期望算法对数据进行填充
            data: 需要填充的数据（一维数组）
            返回值: 填充后的数据（一维数组）
            """
            # 初始化参数
            mean = np.nanmean(data)  # 均值作为初始值
            std = np.nanstd(data)  # 标准差作为初始值
            threshold = 1e-6  # 收敛阈值
            max_iter = 100  # 最大迭代次数

            # 迭代更新参数
            for _ in range(max_iter):
                # E 步：计算每个样本属于正态分布的概率
                if std == 0:
                    break
                prob = np.exp(-0.5 * ((data - mean) / std) ** 2) / (np.sqrt(2 * np.pi) * std)

                # M 步：更新参数
                new_mean = np.nansum(prob * data) / np.nansum(prob)
                new_std = np.sqrt(np.nansum(prob * (data - new_mean) ** 2) / np.nansum(prob))

                # 判断是否收敛
                if np.abs(new_mean - mean) < threshold and np.abs(new_std - std) < threshold:
                    break

                # 更新参数
                mean = new_mean
                std = new_std

            # 处理特殊情况：标准差为零或缺失值
            if std == 0 or np.isnan(std):
                filled_data = np.where(np.isnan(data), np.nanmean(data), data)
            else:
                # 使用估计的参数填充缺失值
                filled_data = np.where(np.isnan(data), np.random.normal(mean, std), data)

            return filled_data

        """
        通过最大期望算法对Excel表中的空白内容进行插补，在指定的范围内执行插补操作。
        filepath:   Excel文件路径
        sheetname:  表名
        left_top:   矩形区域左上角坐标，如 "A1"
        right_bottom: 矩形区域右下角坐标，如 "B1"
        """
        # 将坐标转换为数字索引
        left_top_index = convert_coordinate(left_top)
        right_bottom_index = convert_coordinate(right_bottom)

        # 加载 Excel 文件
        df = pd.read_excel(filepath, sheet_name=sheetname, na_values='', header=None)

        # 获取需要处理的区域
        data_df = df.iloc[left_top_index[0]:right_bottom_index[0], left_top_index[1]:right_bottom_index[1]]

        # 对空白内容的缺失值进行插补
        for col in data_df.columns:
            column_data = data_df[col]

            # 处理数据中的异常值
            column_data = column_data.apply(lambda x: np.nan if x == '' or not np.isfinite(x) else x)

            em_data = em_algorithm(column_data.values)
            data_df[col] = pd.DataFrame(em_data)

        # 将新的数据写回 Excel 文件
        df.iloc[left_top_index[0]:right_bottom_index[0], left_top_index[1]:right_bottom_index[1]] = data_df
        df.to_excel(filepath, sheet_name=sheetname, index=False, header=None)
        flows[active_flow_id].text_pad_1.insert(tk.END, "最大期望法插入成功\n")
    except FileNotFoundError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END,
                          f"最大期望法插入失败:文件路径错误，或不存在对应表格\n",
                          "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except IndexError as e:
        flows[active_flow_id].text_pad_1.insert(tk.END,
                          f"最大期望法插入失败:参数错误或为空\n",
                          "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END,
                          f"最大期望法插入失败{type(e)}\n",
                          "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")


def fill_blank_cells(filepath, sheetname, left_top, right_bottom):
    try:
        # 打开Excel文件
        workbook = openpyxl.load_workbook(filepath)

        # 选择指定的工作表
        worksheet = workbook[sheetname]

        # 获取左上角和右下角的单元格坐标
        left_col, left_row = openpyxl.utils.cell.coordinate_from_string(left_top)
        right_col, right_row = openpyxl.utils.cell.coordinate_from_string(right_bottom)

        # 获取指定范围内的所有单元格
        cells_range = worksheet[left_col:right_col]

        # 存储需要填充的空白单元格的值
        values_to_fill = []

        # 遍历指定范围内的单元格
        for row in cells_range:
            for cell in row:
                # 判断单元格是否为空白
                if cell.value is None or cell.value == "":
                    # 将空白单元格的值添加到列表中
                    values_to_fill.append(None)

        # 检查是否存在需要填充的空白单元格
        if values_to_fill:
            # 去除空值后计算中位数
            values_to_fill = [value for value in values_to_fill if value is not None]
            median_value = statistics.median(values_to_fill)

            # 填充空白单元格
            for row in cells_range:
                for cell in row:
                    if cell.value is None or cell.value == "":
                        cell.value = median_value

            # 保存修改后的Excel文件
            workbook.save(filepath)
            flows[active_flow_id].text_pad_1.insert(tk.END, "中位数填充完成！\n")
        else:
            flows[active_flow_id].text_pad_1.insert(tk.END, "指定范围内没有空白单元格，无需进行填充！\n")
    except openpyxl.utils.exceptions.CellCoordinatesException as cce:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"中位数填充失败：发生单元格坐标错误\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except openpyxl.utils.exceptions.SheetTitleException as ste:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"中位数填充失败：工作表名称错误\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except openpyxl.utils.exceptions.InvalidFileException as fnfe:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"中位数填充失败：参数为空或参数错误\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except statistics.StatisticsError as se:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"中位数填充失败：无法计算中位数\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"中位数填充失败：出现异常{type(e)}\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")


def fill_blank_cells1(filepath, sheetname, left_top, right_bottom):
    try:
        # Open the Excel file
        workbook = openpyxl.load_workbook(filepath)

        # Select the specified worksheet
        worksheet = workbook[sheetname]

        # Get the coordinates of the top-left and bottom-right cells
        left_col, left_row = openpyxl.utils.cell.coordinate_from_string(left_top)
        right_col, right_row = openpyxl.utils.cell.coordinate_from_string(right_bottom)

        # Get all the cells within the specified range
        cells_range = worksheet[left_col:right_col]

        # Store the non-empty values to fill in blank cells
        values_to_fill = []

        # Iterate over the cells within the specified range
        for row in cells_range:
            for cell in row:
                # Check if the cell is not empty
                if cell.value is not None and cell.value != "":
                    # Add the value of non-empty cells to the list
                    values_to_fill.append(cell.value)

        # Check if there are values to calculate the average
        if values_to_fill:
            # Calculate the average
            average_value = statistics.mean(values_to_fill)

            # Fill the blank cells with the average value
            for row in cells_range:
                for cell in row:
                    if cell.value is None or cell.value == "":
                        cell.value = average_value

            # Save the modified Excel file
            workbook.save(filepath)
            flows[active_flow_id].text_pad_1.insert(tk.END, "平均数填充完成！\n")
        else:
            flows[active_flow_id].text_pad_1.insert(tk.END, "指定范围内没有空白单元格，无需进行填充！\n")
    except openpyxl.utils.exceptions.CellCoordinatesException as cce:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"平均数填充失败：发生单元格坐标错误\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except openpyxl.utils.exceptions.SheetTitleException as ste:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"平均数填充失败：工作表名称错误\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except openpyxl.utils.exceptions.InvalidFileException as fnfe:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"平均数填充失败：参数为空或参数错误\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except statistics.StatisticsError as se:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"平均数填充失败：无法计算平均数\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"平均数填充失败：出现异常\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")


def fill_blank_cells2(filepath, sheetname, left_top, right_bottom):
    try:
        # 打开Excel文件
        workbook = openpyxl.load_workbook(filepath)

        # 选择指定的工作表
        worksheet = workbook[sheetname]

        # 获取左上角和右下角的单元格坐标
        left_col, left_row = openpyxl.utils.cell.coordinate_from_string(left_top)
        right_col, right_row = openpyxl.utils.cell.coordinate_from_string(right_bottom)

        # 获取指定范围内的所有单元格
        cells_range = worksheet[left_col:right_col]

        # 存储需要填充的空白单元格的值
        values_to_fill = []

        # 遍历指定范围内的单元格
        for row in cells_range:
            for cell in row:
                # 判断单元格是否为空白
                if cell.value is None or cell.value == "":
                    # 将空白单元格的值添加到列表中
                    values_to_fill.append(None)

        # 检查是否存在需要填充的空白单元格
        if values_to_fill:
            # 计算众数
            mode_value = statistics.mode(values_to_fill)

            # 填充空白单元格
            for row in cells_range:
                for cell in row:
                    if cell.value is None or cell.value == "":
                        cell.value = mode_value

            # 保存修改后的Excel文件
            workbook.save(filepath)

            print("众数填充完成！")
            flows[active_flow_id].text_pad_1.insert(tk.END, "众数填充完成！\n")
        else:
            flows[active_flow_id].text_pad_1.insert(tk.END, "众数填充失败:指定范围内没有空白单元格，无需进行填充！\n")
    except openpyxl.utils.exceptions.CellCoordinatesException as cce:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"众数填充失败:发生单元格坐标错误\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except openpyxl.utils.exceptions.SheetTitleException as ste:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"众数填充失败:工作表名称错误\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except openpyxl.utils.exceptions.InvalidFileException as fnfe:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"众数填充失败:文件不存在\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except statistics.StatisticsError as se:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"众数填充失败:无法计算中位数\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")
    except Exception as e:
        flows[active_flow_id].text_pad_1.insert(tk.END, f"众数填充失败:出现异常\n", "red")
        flows[active_flow_id].text_pad_1.tag_configure("red", foreground="red")


def greater_than(excel_filepath, sheet, search_value, type):
    try:
        df = pd.read_excel(excel_filepath, sheet_name=sheet)

        # 检查是否存在数值类型的 search_value
        if type == "int":
            search_value = int(search_value)
        elif type == "float":
            search_value = float(search_value)
        elif type == "string":
            pass

        # 检查是否存在大于 search_value 的值
        if df[df > search_value].any().any():
            return True

        return False
    except:
        return False


def less_than(excel_filepath, sheet, search_value, type):
    try:
        df = pd.read_excel(excel_filepath, sheet_name=sheet)

        # 检查是否存在数值类型的 search_value
        if type == "int":
            search_value = int(search_value)
        elif type == "float":
            search_value = float(search_value)
        elif type == "string":
            pass

        # 检查是否存在大于 search_value 的值
        if df[df < search_value].any().any():
            return True

        return False
    except:
        return False
def select_window(num):
    windows = {
        1: ExcelFunction1_window,
        2: ExcelFunction2_window,
        3: ExcelFunction3_window,
        4: ExcelFunction4_window,
        6: ExcelFunction6_window,
        7: ExcelFunction7_window,
        8: ExcelFunction8_window,
        9: ExcelFunction9_window,
        10: ExcelFunction10_window,
        11: ExcelFunction11_window,
        12: ExcelFunction12_window,
        13: ExcelFunction13_window,
        14: ExcelFunction14_window,
        15: ExcelFunction15_window,
        16: ExcelFunction16_window,
        17: ExcelFunction17_window,
        18: ExcelFunction18_window,
        19: ExcelFunction19_window,
        20: ExcelFunction20_window,
        100: ExcelFunction100_window
    }

    for n, window_func in windows.items():
        if n == num:
            window_func.place(width=220, height=250)
        else:
            window_func.place_forget()


def excel1(event):
    flows[active_flow_id].text_pad.insert(tk.END, '从mysql读取数据到excel中...\n')
    # Commandlist.bind(replace_excel_values,1)
    flows[active_flow_id].text_pad.edit_separator()
    select_window(1)
def excel2(event):
    flows[active_flow_id].text_pad.insert(tk.END, '从mysql读取数据到word中...\n')
    # flows[active_flow_id].commandcomponent.bind(replace_excel_values,1)
    flows[active_flow_id].text_pad.edit_separator()
    select_window(2)
def excel3(event):
    flows[active_flow_id].text_pad.insert(tk.END, '正在进行替换指定范围内的单元格内容功能...\n')
    # flows[active_flow_id].commandcomponent.bind(replace_excel_values,1)
    flows[active_flow_id].text_pad.edit_separator()
    select_window(3)


def excel4(event):
    flows[active_flow_id].text_pad.insert(tk.END, '正在进行将excel指定范围内的表格插入word文件...\n')
    # flows[active_flow_id].commandcomponent.bind(insert_excel_to_word, 1)
    flows[active_flow_id].text_pad.edit_separator()
    select_window(4)


def excel6(event):
    flows[active_flow_id].text_pad.insert(tk.END, '正在根据表格数据生成多柱状图...\n')
    # flows[active_flow_id].commandcomponent.bind(generate_bar_chart,1)
    flows[active_flow_id].text_pad.edit_separator()
    select_window(6)


def excel7(event):
    flows[active_flow_id].text_pad.insert(tk.END, '正在合并相关文件...\n')
    # flows[active_flow_id].commandcomponent.bind(generate_bar_chart,1)
    flows[active_flow_id].text_pad.edit_separator()
    select_window(7)


def excel8(event):
    flows[active_flow_id].text_pad.insert(tk.END, '正在最大似然法继续填充...\n')
    # flows[active_flow_id].commandcomponent.bind(generate_bar_chart,1)
    flows[active_flow_id].text_pad.edit_separator()
    select_window(8)


def excel9(event):
    flows[active_flow_id].text_pad.insert(tk.END, '正在进行运算填充...\n')
    # flows[active_flow_id].commandcomponent.bind(generate_bar_chart,1)
    flows[active_flow_id].text_pad.edit_separator()
    select_window(9)


def Logical_if1(event):
    flows[active_flow_id].text_pad.insert(tk.END, '判断是否存在对应值\n', "blue_font")
    flows[active_flow_id].text_pad.tag_configure("blue_font", foreground="blue")
    # flows[active_flow_id].commandcomponent.bind(generate_bar_chart,1)
    flows[active_flow_id].text_pad.edit_separator()
    select_window(10)


def Logical_end(event):
    flows[active_flow_id].text_pad.insert(tk.END, '结束逻辑\n', "blue_font")
    flows[active_flow_id].text_pad.tag_configure("blue_font", foreground="blue")
    flows[active_flow_id].commandcomponent.bind(EndCommand())
    select_window(100)
    flows[active_flow_id].text_pad.edit_separator()


def excel11(event):
    flows[active_flow_id].text_pad.insert(tk.END, '从pdf读取数据到word中...\n')
    flows[active_flow_id].text_pad.edit_separator()
    select_window(11)


def excel12(event):
    flows[active_flow_id].text_pad.insert(tk.END, '合并excel中的表格...\n')
    flows[active_flow_id].text_pad.edit_separator()
    select_window(12)

def excel13(event):
    flows[active_flow_id].text_pad.insert(tk.END, '合并word文件...\n')
    flows[active_flow_id].text_pad.edit_separator()
    select_window(13)

def excel14(event):
    flows[active_flow_id].text_pad.insert(tk.END, '合并pdf文件...\n')
    flows[active_flow_id].text_pad.edit_separator()
    select_window(14)


def excel15(event):
    flows[active_flow_id].text_pad.insert(tk.END, '最大期望值法...\n')
    flows[active_flow_id].text_pad.edit_separator()
    select_window(15)

def excel16(event):
    flows[active_flow_id].text_pad.insert(tk.END, '中位数填充...\n')
    flows[active_flow_id].text_pad.edit_separator()
    select_window(16)

def excel17(event):
    flows[active_flow_id].text_pad.insert(tk.END, '平均数填充...\n')
    flows[active_flow_id].text_pad.edit_separator()
    select_window(17)

def excel18(event):
    flows[active_flow_id].text_pad.insert(tk.END, '众数填充...\n')
    flows[active_flow_id].text_pad.edit_separator()
    select_window(18)

def Logical_if2(event):
    flows[active_flow_id].text_pad.insert(tk.END, '判断是否大于对应值\n', "blue_font")
    flows[active_flow_id].text_pad.tag_configure("blue_font", foreground="blue")
    # flows[active_flow_id].commandcomponent.bind(generate_bar_chart,1)
    flows[active_flow_id].text_pad.edit_separator()
    select_window(19)

def Logical_if3(event):
    flows[active_flow_id].text_pad.insert(tk.END, '判断是否小于对应值\n', "blue_font")
    flows[active_flow_id].text_pad.tag_configure("blue_font", foreground="blue")
    # flows[active_flow_id].commandcomponent.bind(generate_bar_chart,1)
    flows[active_flow_id].text_pad.edit_separator()
    select_window(20)
# 创建根窗口
root = tk.Tk()
root.title('表单自动填充软件')
root.geometry('1200x500+100+100')

# 左侧布局
left_name = tk.Frame(root)
left_name.place(x=5, y=5, width=240, height=490)

# 左侧布局 命令列表
list_command_fame = tk.LabelFrame(left_name, text='命令列表', padx=5, pady=5)
list_command_fame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)




# 左侧布局 Excel命令
tree = ttk.Treeview(list_command_fame, height=1)
parent_node = tree.insert("", 0, text="功能列表", open=True)

operation = tree.insert(parent_node, "end", text="规则填充", open=True)
interactive_operation = tree.insert(parent_node, "end", text="交互填充", open=True)
merge_fill = tree.insert(parent_node, "end", text="合并填充111", open=True)
Algorithm_fill = tree.insert(parent_node, "end", text="算法填充", open=True)
logic_command = tree.insert(parent_node, "end", text="条件填充命令", open=True)

export_data_to_excel1 = tree.insert(interactive_operation, "end", text="MYSQL数据填充EXCEL", tags=("1",))
tree.tag_bind("1", "<Button-1>", excel1)
export_data_to_word1 = tree.insert(interactive_operation, "end", text="MYSQL数据填充WORD", tags=("2",))
tree.tag_bind("2", "<Button-1>", excel2)
replace_excel_values1 = tree.insert(operation, "end", text="替换填充", tags=("4",))
tree.tag_bind("4", "<Button-1>", excel3)
insert_excel_to_word1 = tree.insert(interactive_operation, "end", text="EXCEL表格填充WORD", tags=("5",))
tree.tag_bind("5", "<Button-1>", excel4)
generate_bar_chart1 = tree.insert(interactive_operation, "end", text="EXCEL表格数据生成多柱状图", tags=("6",))
tree.tag_bind("6", "<Button-1>", excel6)
merge_Excel_files1 = tree.insert(merge_fill, "end", text="合并EXCEL文件", tags=("7",))
tree.tag_bind("7", "<Button-1>", excel7)
fill_missing_data_in_excel1 = tree.insert(Algorithm_fill, "end", text="最大似然法填充", tags=("8",))
tree.tag_bind("8", "<Button-1>", excel8)
process_excel_data1 = tree.insert(operation, "end", text="运算填充", tags=("9",))
tree.tag_bind("9", "<Button-1>", excel9)
search1= tree.insert(logic_command, "end", text="是否 存在对应值", tags=("10",))
tree.tag_configure("10", foreground="blue")
tree.tag_bind("10", "<Button-1>", Logical_if1)
En1 = tree.insert(logic_command, "end", text="End", tags=("11",))
tree.tag_configure("11", foreground="blue")
tree.tag_bind("11", "<Button-1>", Logical_end)
pdf_convert_word1 = tree.insert(interactive_operation, "end", text="PDF数据填充到WORD", tags=("12",))
tree.tag_bind("12", "<Button-1>", excel11)
merge_excel_sheets1 = tree.insert(merge_fill, "end", text="合并EXCEL表格", tags=("13",))
tree.tag_bind("13", "<Button-1>", excel12)
merge_word_files1 = tree.insert(merge_fill, "end", text="合并WORD文件", tags=("14",))
tree.tag_bind("14", "<Button-1>", excel13)
merge_pdf_files1 = tree.insert(merge_fill, "end", text="合并PDF文件", tags=("15",))
tree.tag_bind("15", "<Button-1>", excel14)
EM_data_in_excel1 = tree.insert(Algorithm_fill, "end", text="最大期望法填充", tags=("16",))
tree.tag_bind("16", "<Button-1>", excel15)
fill_blank_cell1 = tree.insert(Algorithm_fill, "end", text="中位数填充", tags=("17",))
tree.tag_bind("17", "<Button-1>", excel16)
fill_blank_cell2 = tree.insert(Algorithm_fill, "end", text="平均数填充", tags=("18",))
tree.tag_bind("18", "<Button-1>", excel17)
fill_blank_cell3 = tree.insert(Algorithm_fill, "end", text="众数填充", tags=("19",))
tree.tag_bind("19", "<Button-1>", excel18)

greater_than1 = tree.insert(logic_command, "end", text="是否 大于对应值", tags=("20",))
tree.tag_configure("20", foreground="blue")
tree.tag_bind("20", "<Button-1>", Logical_if2)

less_than1 = tree.insert(logic_command, "end", text="是否 小于对应值", tags=("21",))
tree.tag_configure("21", foreground="blue")
tree.tag_bind("21", "<Button-1>", Logical_if3)


class ToolTip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tip_window = None

    def show(self, x, y):
        if self.tip_window or not self.text:
            return
        self.tip_window = tk.Toplevel(self.widget)
        self.tip_window.wm_overrideredirect(1)
        self.tip_window.wm_geometry(f"+{x}+{y}")

        label = tk.Label(
            self.tip_window,
            text=self.text,
            background="white",
            relief=tk.SOLID,
            borderwidth=0,
            font=('楷体', 10, "normal"),  # 设置字体样式和大小
            foreground="red",  # 设置字体颜色
            anchor="center" ,
            justify="center",
            wraplength = 400
        )
        label.pack()

    def hide(self):
        if self.tip_window:
            self.tip_window.destroy()
            self.tip_window = None


def on_item_hover(event):
    item_id = tree.identify("item", event.x, event.y)
    tooltip = tooltips.get(item_id)
    if tooltip:
        for tip in tooltips.values():
            if tip != tooltip:
                tip.hide()
        tooltip.show(event.x_root + 10, event.y_root + 10)

# 鼠标离开事件处理函数
def on_item_leave(event):
    for tooltip in tooltips.values():
        tooltip.hide()

tree.bind("<Motion>", on_item_hover)
tree.bind("<Leave>", on_item_leave)

tooltips = {
    export_data_to_excel1: ToolTip(tree, "功能：将指定的MYSQL数据库的表其中的内容导入到指定路径指定表格名称的EXCEL文件中 参数:host:主机名:string, port:端口号:int, user:用户名:string, password:用户密码:string, database:数据库名称:string, table:数据表名称:string, output_file:保存的文件绝对路径:string"),
    export_data_to_word1: ToolTip(tree,"功能将指定的MYSQL数据库的表其中的内容导入到指定路径指定表格名称的WORD文件中 参数:host:主机名:string, port:端口号:int, user:用户名:string, password:用户密码:string, database:数据库名称:string, table:数据表名称:string, output_file:保存的文件绝对路径:string"),
    replace_excel_values1: ToolTip(tree,"功能:将指定EXCEL文件中指定表格、指定位置范围内的某个特定数据进行替换 参数：filepath：文件路径:string ,sheet_name:表格名称:string start_cell:范围左上角坐标:string，例如 'A1',end_cell:范围右下角角坐标:string，例如 'H10',old_value:被替换的数据:object , new_value: 替换后的数据:object "),
    insert_excel_to_word1: ToolTip(tree,"功能:将指定EXCEL文件中指定表格插入到指定WORD文件 参数：filepath：EXCEL文件路径:string ,sheet_name:EXCEL选取表格名称:string ,word_filepath：EXCEL文件路径:string ,start_coordinate:范围左上角坐标:string，例如 'A1' ,end_coordinate:范围右下角角坐标:string，例如 'A1'"),
    generate_bar_chart1: ToolTip(tree,"功能:将指定范围内的EXCEL表格内的数据，以左上角坐标所在行为数值列，以所在列为行索引(都不包括左上角坐标本身)生成一个jpg格式的文件 参数：filepath：EXCEL文件路径:string ,sheet_name:EXCEL选取表格名称:string ,word_filepath：EXCEL文件路径:string ,start_coordinate:范围左上角坐标:string，例如 'A1' ,end_coordinate:范围右下角角坐标:string，例如 'A1',picture_path:图片的保存路径:string"),
    merge_Excel_files1: ToolTip(tree,"功能:将多个EXCEL文件合并到一个指定路径的EXCEL文件中,参数:Summary_document_path:合并后的文件路径:string ,*file_paths:需要合并的文件列表:List"),
    fill_missing_data_in_excel1: ToolTip(tree,"功能:用最大似然法对指定路径指定表格的excel文件指定范围的空白采用最大释然法进行填充,参数:filepath：文件路径:string ,sheet_name:表格名称:string left_top:范围左上角坐标:string，例如 'A1',right_bottom:范围右下角角坐标:string，例如 'H10'"),
    process_excel_data1: ToolTip(tree,"功能:将被操作的列的内容乘以权重和操作列号的内容进行对应操作后存储到操作列号中, 参数:file_path：EXCEL文件路径:string , sheet_name:EXCEL选取表格名称:string ,column_name1:被操作的列号:string 例如'A', weight:被操作的列号的权重:float , column_name2:操作的列号:string 例如'A', operation:操作的方法:string 例如'+'"),
    search1:ToolTip(tree,"功能:查找指对应的EXCEL文件的对应表格中是否存在对应类型的值, 参数:excel_filepath:EXCEL文件路径:string , sheet:EXCEL选取表格名称:string, search_value:查询的值:type,type:值的类型:string 返回值:TRUE OR FALSE"),
    En1:ToolTip(tree, "功能:用于结束条件逻辑, 参数:无"),
    pdf_convert_word1: ToolTip(tree,"功能:将指定范围内的pdf文件内容提取到WORD文件中:参数:pdf_filepath:PDF文件路径:string , start_page:PDF文件起始页:string , end_page:PDF文件结束页:string ,output_file:WORD文件输出路径:string"),
    merge_excel_sheets1: ToolTip(tree,"功能:将两个EXCEL文件的表格合并后填充到指定的EXCEL文件的表格中,参数:excel_file1_path:EXCEL1文件路径:string , excel_file2_path:EXCEL1文件路径:string , sheet_name1:EXCEL1文件的表格名称:string , sheet_name2:EXCEL1文件的表格名称:string , merged_file_path:合并的EXCEL文件路径:string , merged_sheet_name合并的EXCEL文件表格名称:string "),
    merge_word_files1: ToolTip(tree,"功能:将两个WORD文件的指定范围的内容合并到指定路径的WORD中, 参数:file1_path:WORD1文件路径:string , file1_start:WORD1文件起始段落号:string, file1_end:WORD1文件结束段落号:string, file2_path:WORD2文件路径:string, file2_start:WORD2文件起始段落号:string, file2_end:WORD2文件结束段落号:string, output_path:合并后的文件路径:string"),
    merge_pdf_files1: ToolTip(tree,"功能:将两个PDF文件的指定范围的内容合并到指定路径的PDF中, 参数:file1_path:PDF1文件路径:string , file1_start:PDF1文件起始段落号:string, file1_end:PDF1文件结束段落号:string, file2_path:PDF2文件路径:string, file2_start:PDF2文件起始段落号:string, file2_end:PDF2文件结束段落号:string, output_path:合并后的文件路径:string"),
    EM_data_in_excel1: ToolTip(tree,"功能:用最大期望算法对指定路径指定表格的excel文件指定范围的空白进行填充,参数:filepath：文件路径:string ,sheet_name:表格名称:string left_top:范围左上角坐标:string，例如 'A1',right_bottom:范围右下角角坐标:string，例如 'H10'"),
    fill_blank_cell1: ToolTip(tree,"功能:用中位数对指定路径指定表格的excel文件指定范围的空白进行填充,参数:filepath：文件路径:string ,sheet_name:表格名称:string left_top:范围左上角坐标:string，例如 'A1',right_bottom:范围右下角角坐标:string，例如 'H10'"),
    fill_blank_cell2: ToolTip(tree, "功能:用平均数对指定路径指定表格的excel文件指定范围的空白进行填充,参数:filepath：文件路径:string ,sheet_name:表格名称:string left_top:范围左上角坐标:string，例如 'A1',right_bottom:范围右下角角坐标:string，例如 'H10'"),
    fill_blank_cell3: ToolTip(tree,"功能:用众数对指定路径指定表格的excel文件指定范围的空白进行填充,参数:filepath：文件路径:string ,sheet_name:表格名称:string left_top:范围左上角坐标:string，例如 'A1',right_bottom:范围右下角角坐标:string，例如 'H10'"),
    greater_than1: ToolTip(tree,"功能:查找指对应的EXCEL文件的对应表格中是否存在大于对应类型的值,参数:excel_filepath:EXCEL文件路径:string , sheet:EXCEL选取表格名称:string, search_value:查询的值:type,type:值的类型:string 返回值:TRUE OR FALSE"),
    less_than1: ToolTip(tree,"功能:查找指对应的EXCEL文件的对应表格中是否存在小于对应类型的值,参数:excel_filepath:EXCEL文件路径:string , sheet:EXCEL选取表格名称:string, search_value:查询的值:type,type:值的类型:string 返回值:TRUE OR FALSE")
}


tree.pack(side="left", fill="both", expand=True)

# 中间布局
middle_frame = tk.Frame(root)
middle_frame.place(x=250, y=5, width=590, height=490)

info_frame = tk.LabelFrame(middle_frame, text='命令添加成功信息提示窗口')
info_frame.place(x=5, y=5, width=575, height=370)

"""text_pad = tk.Text(info_frame, width=58, height=25, autoseparators=False, undo=True, maxundo=100)
text_pad.pack(fill=tk.BOTH, expand=True)"""
sendMessage_frame = tk.LabelFrame(middle_frame, text='运行信息提示窗口')
sendMessage_frame.place(x=5, y=380, width=575, height=110)

sendMessage_frame_Children = tk.Frame(sendMessage_frame)
sendMessage_frame_Children.pack(fill=tk.BOTH, expand=True)

"""text_pad_1 = tk.Text(sendMessage_frame_Children, width=58, height=6)
text_pad_1.pack(fill=tk.BOTH, expand=True)"""

# flows[active_flow_id].commandcomponent = CommandComponent()

# 右侧布局
right_name = tk.Frame(root)
right_name.place(x=850, y=5, width=340, height=500)

info_frame1 = tk.LabelFrame(right_name, text='参数')
info_frame1.place(x=10, y=10, width=320, height=250)


# host, port, user, password, database, table, output_file
ExcelFunction100_window = tk.Frame(info_frame1)
"""----------Excel-export_data_to_excel------------"""
ExcelFunction1_window = tk.Frame(info_frame1)
host = tk.Label(ExcelFunction1_window, text='主机地址', font=('Arial', 12),justify='left', anchor='w')
host.place(x=10, y=10)
entry_path_host = ttk.Entry(ExcelFunction1_window, font=('Arial', 12))
entry_path_host.place(x=95, y=10, width=115)

port = tk.Label(ExcelFunction1_window, text='端口号', font=('Arial', 12),justify='left', anchor='w')
port.place(x=10, y=35)
entry_path_port = ttk.Entry(ExcelFunction1_window, font=('Arial', 12))
entry_path_port.place(x=95, y=35, width=115)

user = tk.Label(ExcelFunction1_window, text='用户名', font=('Arial', 12),justify='left', anchor='w')
user.place(x=10, y=60)
entry_path_start_user= ttk.Entry(ExcelFunction1_window, font=('Arial', 12))
entry_path_start_user.place(x=95, y=60, width=115)

password = tk.Label(ExcelFunction1_window, text='密码', font=('Arial', 12),justify='left', anchor='w')
password.place(x=10, y=85)
entry_path_end_password = ttk.Entry(ExcelFunction1_window, font=('Arial', 12))
entry_path_end_password.place(x=95, y=85, width=115)

database = tk.Label(ExcelFunction1_window, text='数据库名', font=('Arial', 12),justify='left', anchor='w')
database.place(x=10, y=110)
entry_path_database = ttk.Entry(ExcelFunction1_window, font=('Arial', 12))
entry_path_database.place(x=95, y=110, width=115)

table= tk.Label(ExcelFunction1_window, text='数据表名', font=('Arial', 12),justify='left', anchor='w')
table.place(x=10, y=135)
entry_path_table = ttk.Entry(ExcelFunction1_window, font=('Arial', 12))
entry_path_table.place(x=95, y=135, width=115)

output_file= tk.Label(ExcelFunction1_window, text='excel文件路径', font=('Arial', 12),justify='left', anchor='w')
output_file.place(x=10, y=160)
entry_path_output_file = ttk.Entry(ExcelFunction1_window, font=('Arial', 12))
entry_path_output_file.place(x=95, y=160, width=115)

button_save_e1 = tk.Button(ExcelFunction1_window, text='添加', font=('Arial', 12))
button_save_e1.place(x=70, y=205, width=90)
button_save_e1.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(export_data_to_excel,1,entry_path_host.get().strip(),
                                                     entry_path_port.get().strip(),
                                                     entry_path_start_user.get().strip(),
                                                     entry_path_end_password.get().strip(),
                                                     entry_path_database.get().strip(),
                                                     entry_path_table.get().strip(),
                                                     entry_path_output_file.get().strip()
                                                    )
                                     )
)
"""-------------------------------------------"""
"""----------Excel-export_data_to_word------------"""
ExcelFunction2_window = tk.Frame(info_frame1)
host1 = tk.Label(ExcelFunction2_window, text='主机地址', font=('Arial', 12),justify='left', anchor='w')
host1.place(x=10, y=10)
entry_path_host1 = ttk.Entry(ExcelFunction2_window, font=('Arial', 12))
entry_path_host1.place(x=95, y=10, width=115)

port1 = tk.Label(ExcelFunction2_window, text='端口号', font=('Arial', 12),justify='left', anchor='w')
port1.place(x=10, y=35)
entry_path_port1 = ttk.Entry(ExcelFunction2_window, font=('Arial', 12))
entry_path_port1.place(x=95, y=35, width=115)

user1 = tk.Label(ExcelFunction2_window, text='用户名', font=('Arial', 12),justify='left', anchor='w')
user1.place(x=10, y=60)
entry_path_start_user1= ttk.Entry(ExcelFunction2_window, font=('Arial', 12))
entry_path_start_user1.place(x=95, y=60, width=115)

password1 = tk.Label(ExcelFunction2_window, text='密码', font=('Arial', 12),justify='left', anchor='w')
password1.place(x=10, y=85)
entry_path_end_password1 = ttk.Entry(ExcelFunction2_window, font=('Arial', 12))
entry_path_end_password1.place(x=95, y=85, width=115)

database1 = tk.Label(ExcelFunction2_window, text='数据库名', font=('Arial', 12),justify='left', anchor='w')
database1.place(x=10, y=110)
entry_path_database1 = ttk.Entry(ExcelFunction2_window, font=('Arial', 12))
entry_path_database1.place(x=95, y=110, width=115)

table1= tk.Label(ExcelFunction2_window, text='数据表名', font=('Arial', 12),justify='left', anchor='w')
table1.place(x=10, y=135)
entry_path_table1 = ttk.Entry(ExcelFunction2_window, font=('Arial', 12))
entry_path_table1.place(x=95, y=135, width=115)

output_file1= tk.Label(ExcelFunction2_window, text='word文件路径', font=('Arial', 12),justify='left', anchor='w')
output_file1.place(x=10, y=135)
entry_path_output_file1 = ttk.Entry(ExcelFunction2_window, font=('Arial', 12))
entry_path_output_file1.place(x=95, y=135, width=115)

button_save_e2 = tk.Button(ExcelFunction2_window, text='添加', font=('Arial', 12))
button_save_e2.place(x=70, y=180, width=90)
button_save_e2.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(export_data_to_word,2,entry_path_host1.get().strip(),
                                                     entry_path_port1.get().strip(),
                                                     entry_path_start_user1.get().strip(),
                                                     entry_path_end_password1.get().strip(),
                                                     entry_path_database1.get().strip(),
                                                     entry_path_table1.get().strip(),
                                                     entry_path_output_file1.get().strip()
                                                    )
                                     )
)
"""-------------------------------------------"""
"""----------Excel-replace_excel_values------------"""
ExcelFunction3_window = tk.Frame(info_frame1)
filepath3 = tk.Label(ExcelFunction3_window, text='Excel文件路径', font=('Arial', 12),justify='left', anchor='w')
filepath3.place(x=10, y=10)
entry_path_filepath3 = ttk.Entry(ExcelFunction3_window, font=('Arial', 12))
entry_path_filepath3.place(x=95, y=10, width=115)

sheet_name3 = tk.Label(ExcelFunction3_window, text='表格名称', font=('Arial', 12),justify='left', anchor='w')
sheet_name3.place(x=10, y=35)
entry_path_sheet_name3 = ttk.Entry(ExcelFunction3_window, font=('Arial', 12))
entry_path_sheet_name3.place(x=95, y=35, width=115)

start_cell_name = tk.Label(ExcelFunction3_window, text='左上角坐标', font=('Arial', 12),justify='left', anchor='w')
start_cell_name.place(x=10, y=60)
entry_path_start_cell_name = ttk.Entry(ExcelFunction3_window, font=('Arial', 12))
entry_path_start_cell_name.place(x=95, y=60, width=115)

end_cell_name = tk.Label(ExcelFunction3_window, text='右下角坐标', font=('Arial', 12),justify='left', anchor='w')
end_cell_name.place(x=10, y=85)
entry_path_end_cell_name = ttk.Entry(ExcelFunction3_window, font=('Arial', 12))
entry_path_end_cell_name.place(x=95, y=85, width=115)

old_value_name = tk.Label(ExcelFunction3_window, text='被替换的值', font=('Arial', 12),justify='left', anchor='w')
old_value_name.place(x=10, y=110)
entry_path_old_value_name = ttk.Entry(ExcelFunction3_window, font=('Arial', 12))
entry_path_old_value_name.place(x=95, y=110, width=115)

new_value_name = tk.Label(ExcelFunction3_window, text='替换后的值', font=('Arial', 12),justify='left', anchor='w')
new_value_name.place(x=10, y=135)
entry_path_new_value_name = ttk.Entry(ExcelFunction3_window, font=('Arial', 12))
entry_path_new_value_name.place(x=95, y=135, width=115)

button_save_e3 = tk.Button(ExcelFunction3_window, text='添加', font=('Arial', 12))
button_save_e3.place(x=70, y=180, width=90)
button_save_e3.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(replace_excel_values,3,entry_path_filepath3.get().strip(),
                                                     entry_path_sheet_name3.get().strip(),
                                                     entry_path_start_cell_name.get().strip(),
                                                     entry_path_end_cell_name.get().strip(),
                                                     entry_path_old_value_name.get().strip(),
                                                     entry_path_new_value_name.get().strip()
                                                    )
                                     )
)
"""-------------------------------------------"""


"""----------Excel-insert_excel_to_word------------"""

ExcelFunction4_window = tk.Frame(info_frame1)

filepath4 = tk.Label(ExcelFunction4_window, text='Excel文件路径', font=('Arial', 12))
filepath4.place(x=10, y=10)
entry_path_filepath4 = ttk.Entry(ExcelFunction4_window, font=('Arial', 12))
entry_path_filepath4.place(x=95, y=10, width=115)
sheet_name4 = tk.Label(ExcelFunction4_window, text='表格名称', font=('Arial', 12))
sheet_name4.place(x=10, y=35)
entry_path_sheet_name4 = ttk.Entry(ExcelFunction4_window, font=('Arial', 12))
entry_path_sheet_name4.place(x=95, y=35, width=115)

word_filepath = tk.Label(ExcelFunction4_window, text='word文件路径', font=('Arial', 12))
word_filepath.place(x=10, y=60)
entry_path_word_filepath = ttk.Entry(ExcelFunction4_window, font=('Arial', 12))
entry_path_word_filepath.place(x=95, y=60, width=115)
start_coordinate = tk.Label(ExcelFunction4_window, text='左上角坐标', font=('Arial', 12))
start_coordinate.place(x=10, y=85)
entry_path_start_coordinate = ttk.Entry(ExcelFunction4_window, font=('Arial', 12))
entry_path_start_coordinate.place(x=95, y=85, width=115)
end_coordinate = tk.Label(ExcelFunction4_window, text='右下角坐标', font=('Arial', 12))
end_coordinate.place(x=10, y=110)
entry_path_end_coordinate = ttk.Entry(ExcelFunction4_window, font=('Arial', 12))
entry_path_end_coordinate.place(x=95, y=110, width=115)
button_save_e4 = tk.Button(ExcelFunction4_window, text='添加', font=('Arial', 12))
button_save_e4.place(x=70, y=155, width=90)
button_save_e4.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(insert_excel_to_word, 4,
                                             entry_path_filepath4.get().strip(),
                                             entry_path_sheet_name4.get().strip(),
                                             entry_path_word_filepath.get().strip(),
                                             entry_path_start_coordinate.get().strip(),
                                             entry_path_end_coordinate.get().strip())))

"""-------------------------------------------"""


"""----------Excel-generate_bar_chart------------"""

ExcelFunction6_window = tk.Frame(info_frame1)

filepath6 = tk.Label(ExcelFunction6_window, text='Excel文件路径', font=('Arial', 12))
filepath6.place(x=10, y=10)
entry_path_filepath6 = ttk.Entry(ExcelFunction6_window, font=('Arial', 12))
entry_path_filepath6.place(x=95, y=10, width=115)
sheet_name6 = tk.Label(ExcelFunction6_window, text='表格名称', font=('Arial', 12))
sheet_name6.place(x=10, y=35)
entry_path_sheet_name6 = ttk.Entry(ExcelFunction6_window, font=('Arial', 12))
entry_path_sheet_name6.place(x=95, y=35, width=115)

top_left_coord = tk.Label(ExcelFunction6_window, text='左上角坐标:', font=('Arial', 12))
top_left_coord.place(x=10, y=60)
entry_path_top_left_coord = ttk.Entry(ExcelFunction6_window, font=('Arial', 12))
entry_path_top_left_coord.place(x=95, y=60, width=115)
bottom_right_coord = tk.Label(ExcelFunction6_window, text='右下角坐标:', font=('Arial', 12))
bottom_right_coord.place(x=10, y=85)
entry_path_bottom_right_coord = ttk.Entry(ExcelFunction6_window, font=('Arial', 12))
entry_path_bottom_right_coord.place(x=95, y=85, width=115)
picture_path = tk.Label(ExcelFunction6_window, text='图片保存路径:', font=('Arial', 12))
picture_path.place(x=10, y=110)
entry_path_picture_path = ttk.Entry(ExcelFunction6_window, font=('Arial', 12))
entry_path_picture_path.place(x=95, y=110, width=115)
button_save_e4 = tk.Button(ExcelFunction6_window, text='添加', font=('Arial', 12))
button_save_e4.place(x=70, y=160, width=90)
button_save_e4.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(generate_bar_chart, 6,
                                             entry_path_filepath6.get().strip(),
                                             entry_path_sheet_name6.get().strip(),
                                             entry_path_top_left_coord.get().strip(),
                                             entry_path_bottom_right_coord.get().strip(),
                                             entry_path_picture_path.get().strip())))
"""-------------------------------------------"""


"""----------Excel-merge_Excel_files------------"""
def select_excel_files():
    excel_files = filedialog.askopenfilenames(filetypes=[("Excel Files", "*.xlsx")])
    return excel_files
ExcelFunction7_window = tk.Frame(info_frame1)
merge_Excel_files_path = tk.Label(ExcelFunction7_window, text='合并后的文件保存路径:', font=('Arial', 12))
merge_Excel_files_path.place(x=45, y=10)
entry_path_merge_Excel_files_path = ttk.Entry(ExcelFunction7_window, font=('Arial', 12))
entry_path_merge_Excel_files_path.place(x=45, y=30, width=135)
def save_selected_files():
    selected_files = select_excel_files()
    flows[active_flow_id].commandcomponent.bind(OperateCommand(merge_Excel_files, 7, entry_path_merge_Excel_files_path.get().strip(), *selected_files))
button_save_e5 = tk.Button(ExcelFunction7_window, text='选取并添加', font=('Arial', 12))
button_save_e5.place(x=45, y=105, width=135)
button_save_e5.config(
    command=save_selected_files)
"""-------------------------------------------"""


"""----------Excel-fill_missing_data_in_excel(最大似然法填充表格)------------"""
ExcelFunction8_window = tk.Frame(info_frame1)



fill_missing_data_in_excel_filepath_label = tk.Label(ExcelFunction8_window, text='excel文件路径', font=('Arial', 12))
fill_missing_data_in_excel_filepath_label.place(x=10, y=10)
fill_missing_data_in_excel_filepath1 = ttk.Entry(ExcelFunction8_window, font=('Arial', 12))
fill_missing_data_in_excel_filepath1.place(x=95, y=10, width=115)

fill_missing_data_in_excel_sheetname_label = tk.Label(ExcelFunction8_window, text='表格名称', font=('Arial', 12))
fill_missing_data_in_excel_sheetname_label.place(x=10, y=35)
fill_missing_data_in_excel_sheetname1 = ttk.Entry(ExcelFunction8_window, font=('Arial', 12))
fill_missing_data_in_excel_sheetname1.place(x=95, y=35, width=115)

left_top_coord = tk.Label(ExcelFunction8_window, text='左上角角坐标:', font=('Arial', 12))
left_top_coord.place(x=10, y=60)
entry_path_left_top_coord1 = ttk.Entry(ExcelFunction8_window, font=('Arial', 12))
entry_path_left_top_coord1.place(x=95, y=60, width=115)

right_bottom_coord = tk.Label(ExcelFunction8_window, text='右下角坐标:', font=('Arial', 12))
right_bottom_coord.place(x=10, y=85)
entry_path_right_bottom_coord1 = ttk.Entry(ExcelFunction8_window, font=('Arial', 12))
entry_path_right_bottom_coord1.place(x=95, y=85, width=115)

button_save_e110 = tk.Button(ExcelFunction8_window, text='添加', font=('Arial', 12))
button_save_e110.place(x=70, y=130, width=90)
button_save_e110.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(fill_missing_data_in_excel, 8,
                                                    fill_missing_data_in_excel_filepath1.get().strip(),
                                                    fill_missing_data_in_excel_sheetname1.get().strip(),
                                                     entry_path_left_top_coord1.get().strip(),
                                                     entry_path_right_bottom_coord1.get().strip())))
"""-------------------------------------------"""


"""----------Excel-if------------"""
ExcelFunction10_window = tk.Frame(info_frame1)
if_excel_filepath_label = tk.Label(ExcelFunction10_window, text='excel文件路径', font=('Arial', 12))
if_excel_filepath_label.place(x=10, y=10)
if_excel_filepath10 = ttk.Entry(ExcelFunction10_window, font=('Arial', 12))
if_excel_filepath10.place(x=95, y=10, width=115)

if_excel_sheetname_label = tk.Label(ExcelFunction10_window, text='表格名称', font=('Arial', 12))
if_excel_sheetname_label.place(x=10, y=35)
if_excel_sheetname10 = ttk.Entry(ExcelFunction10_window, font=('Arial', 12))
if_excel_sheetname10.place(x=95, y=35, width=115)

search_value = tk.Label(ExcelFunction10_window, text='检索值:', font=('Arial', 12))
search_value.place(x=10, y=60)
entry_path_search_value10 = ttk.Entry(ExcelFunction10_window, font=('Arial', 12))
entry_path_search_value10.place(x=95, y=60, width=115)


type_value = tk.Label(ExcelFunction10_window, text='检索值类型:', font=('Arial', 12))
type_value.place(x=10, y=85)
type_value10 = ttk.Combobox(ExcelFunction10_window)
type_value10.place(x=95, y=85, width=115)
type_value10['values'] = ('int', 'float', 'string')
button_save_e10 = tk.Button(ExcelFunction10_window, text='添加', font=('Arial', 12))
button_save_e10.place(x=70, y=140, width=90)
button_save_e10.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(LogicalCommand(search,10,if_excel_filepath10.get().strip(),if_excel_sheetname10.get().strip(),
                                                    entry_path_search_value10.get().strip(),type_value10.get().strip())))
"""-------------------------------------------"""


"""----------Excel-process_excel_data(运算填充表格)------------"""
ExcelFunction9_window = tk.Frame(info_frame1)
process_excel_data_filepath_label = tk.Label(ExcelFunction9_window, text='excel文件路径', font=('Arial', 12))
process_excel_data_filepath_label.place(x=10, y=10)
process_excel_data_filepath = ttk.Entry(ExcelFunction9_window, font=('Arial', 12))
process_excel_data_filepath.place(x=95, y=10, width=115)

process_excel_data_sheetname_label = tk.Label(ExcelFunction9_window, text='表格名称', font=('Arial', 12))
process_excel_data_sheetname_label.place(x=10, y=35)
process_excel_data_sheetname = ttk.Entry(ExcelFunction9_window, font=('Arial', 12))
process_excel_data_sheetname.place(x=95, y=35, width=115)

column_name1 = tk.Label(ExcelFunction9_window, text='被计算列名:', font=('Arial', 12))
column_name1.place(x=10, y=60)
entry_path_column_name1= ttk.Entry(ExcelFunction9_window, font=('Arial', 12))
entry_path_column_name1.place(x=95, y=60, width=115)

weight = tk.Label(ExcelFunction9_window, text='权重:', font=('Arial', 12))
weight.place(x=10, y=85)
entry_path_weight = ttk.Entry(ExcelFunction9_window, font=('Arial', 12))
entry_path_weight.place(x=95, y=85, width=115)

column_name2 = tk.Label(ExcelFunction9_window, text='计算列名:', font=('Arial', 12))
column_name2.place(x=10, y=110)
entry_path_column_name2= ttk.Entry(ExcelFunction9_window, font=('Arial', 12))
entry_path_column_name2.place(x=95, y=110, width=115)
operation = tk.Label(ExcelFunction9_window, text='操作类型:', font=('Arial', 12))
operation.place(x=10, y=135)
operation_box = ttk.Combobox(ExcelFunction9_window)
operation_box.place(x=95, y=135, width=115)
operation_box['values'] = ('+', '-', '*', '/')


button_save_e11 = tk.Button(ExcelFunction9_window, text='添加', font=('Arial', 12))
button_save_e11 .place(x=70, y=170, width=90)
button_save_e11 .config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(process_excel_data, 9,process_excel_data_filepath.get().strip(),
                                     process_excel_data_sheetname.get().strip(),entry_path_column_name1.get().strip(),
                                     (entry_path_weight.get().strip()),entry_path_column_name2.get().strip(),
                                     operation_box.get().strip())))


"""----------Excel-pdf_convert_word------------"""
ExcelFunction11_window = tk.Frame(info_frame1)
pdf_filepath_label = tk.Label(ExcelFunction11_window, text='pdf文件路径', font=('Arial', 12))
pdf_filepath_label.place(x=10, y=10)
entry_pdf_filepath = ttk.Entry(ExcelFunction11_window, font=('Arial', 12))
entry_pdf_filepath.place(x=95, y=10, width=115)

start_page_label = tk.Label(ExcelFunction11_window, text='pdf起始页号', font=('Arial', 12))
start_page_label.place(x=10, y=35)
entry_start_page = ttk.Entry(ExcelFunction11_window, font=('Arial', 12))
entry_start_page.place(x=95, y=35, width=115)

end_page_label = tk.Label(ExcelFunction11_window, text='pdf结束页号', font=('Arial', 12))
end_page_label .place(x=10, y=60)
entry_end_page= ttk.Entry(ExcelFunction11_window, font=('Arial', 12))
entry_end_page.place(x=95, y=60, width=115)

output_file_label = tk.Label(ExcelFunction11_window, text='word文件路径', font=('Arial', 12))
output_file_label.place(x=10, y=85)
entry_output_file = ttk.Entry(ExcelFunction11_window, font=('Arial', 12))
entry_output_file.place(x=95, y=85, width=115)

button_save_e12 = tk.Button(ExcelFunction11_window, text='添加', font=('Arial', 12))
button_save_e12.place(x=70, y=130, width=90)
button_save_e12.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(pdf_convert_word, 11,entry_pdf_filepath.get().strip(),entry_start_page.get().strip(),
                                                    entry_end_page.get().strip(),entry_output_file.get().strip())))

# excel_file1_path, excel_file2_path, sheet_name1, sheet_name2, merged_file_path, merged_sheet_name

"""----------Excel-merge_excel_sheets------------"""
ExcelFunction12_window = tk.Frame(info_frame1)
excel_file1_path_label = tk.Label(ExcelFunction12_window, text='excel1文件路径', font=('Arial', 12))
excel_file1_path_label.place(x=10, y=10)
process_excel_file1_path = ttk.Entry(ExcelFunction12_window, font=('Arial', 12))
process_excel_file1_path.place(x=105, y=10, width=115)

excel_file2_path_label = tk.Label(ExcelFunction12_window, text='excel2文件路径', font=('Arial', 12))
excel_file2_path_label.place(x=10, y=35)
process_excel_file2_path = ttk.Entry(ExcelFunction12_window, font=('Arial', 12))
process_excel_file2_path.place(x=105, y=35, width=115)

sheet_name1 = tk.Label(ExcelFunction12_window, text='excel1的表名', font=('Arial', 12))
sheet_name1.place(x=10, y=60)
entry_path_sheet_name1= ttk.Entry(ExcelFunction12_window, font=('Arial', 12))
entry_path_sheet_name1.place(x=105, y=60, width=115)

sheet_name2 = tk.Label(ExcelFunction12_window, text='excel2的表名', font=('Arial', 12))
sheet_name2.place(x=10, y=85)
entry_path_sheet_name2 = ttk.Entry(ExcelFunction12_window, font=('Arial', 12))
entry_path_sheet_name2.place(x=105, y=85, width=115)

merged_file_path = tk.Label(ExcelFunction12_window, text='合并后文件路径', font=('Arial', 12))
merged_file_path.place(x=10, y=110)
entry_path_merged_file_path= ttk.Entry(ExcelFunction12_window, font=('Arial', 12))
entry_path_merged_file_path.place(x=105, y=110, width=115)

merged_sheet_name= tk.Label(ExcelFunction12_window, text='合并后表格名称', font=('Arial', 12))
merged_sheet_name.place(x=10, y=135)
entry_path_merged_sheet_name = ttk.Entry(ExcelFunction12_window, font=('Arial', 12))
entry_path_merged_sheet_name.place(x=105, y=135, width=115)

button_save_e13 = tk.Button(ExcelFunction12_window, text='添加', font=('Arial', 12))
button_save_e13.place(x=70, y=180, width=90)
button_save_e13.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(merge_excel_sheets, 12,
process_excel_file1_path.get().strip(),process_excel_file2_path.get().strip(),entry_path_sheet_name1.get().strip(),
entry_path_sheet_name2.get().strip(),entry_path_merged_file_path.get().strip(),entry_path_merged_sheet_name.get().strip()
                                                    )))

"""----------Excel-merge_word_files------------"""
#  file1_path, file1_start, file1_end, file2_path, file2_start, file2_end, output_path):

ExcelFunction13_window = tk.Frame(info_frame1)
file1_path_label = tk.Label(ExcelFunction13_window, text='word1文件路径', font=('Arial', 12))
file1_path_label.place(x=10, y=10)
process_file1_path= ttk.Entry(ExcelFunction13_window, font=('Arial', 12))
process_file1_path.place(x=105, y=10, width=115)

file1_start_label = tk.Label(ExcelFunction13_window, text='word1的起始页', font=('Arial', 12))
file1_start_label.place(x=10, y=35)
process_file1_start = ttk.Entry(ExcelFunction13_window, font=('Arial', 12))
process_file1_start.place(x=105, y=35, width=115)

file1_end = tk.Label(ExcelFunction13_window, text='word1的总结页', font=('Arial', 12))
file1_end.place(x=10, y=60)
entry_path_file1_end= ttk.Entry(ExcelFunction13_window, font=('Arial', 12))
entry_path_file1_end.place(x=105, y=60, width=115)

file2_path = tk.Label(ExcelFunction13_window, text='word2文件路径', font=('Arial', 12))
file2_path.place(x=10, y=85)
entry_path_file2_path= ttk.Entry(ExcelFunction13_window, font=('Arial', 12))
entry_path_file2_path.place(x=105, y=85, width=115)

file2_start = tk.Label(ExcelFunction13_window, text='word2的起始页', font=('Arial', 12))
file2_start.place(x=10, y=110)
entry_path_file2_start = ttk.Entry(ExcelFunction13_window, font=('Arial', 12))
entry_path_file2_start.place(x=105, y=110, width=115)

file2_end = tk.Label(ExcelFunction13_window, text='word2的总结页', font=('Arial', 12))
file2_end.place(x=10, y=135)
entry_path_file2_end = ttk.Entry(ExcelFunction13_window, font=('Arial', 12))
entry_path_file2_end.place(x=105, y=135, width=115)

output_path = tk.Label(ExcelFunction13_window, text='合并后文件路径', font=('Arial', 12))
output_path.place(x=10, y=160)
entry_path_output_path= ttk.Entry(ExcelFunction13_window, font=('Arial', 12))
entry_path_output_path.place(x=105, y=160, width=115)

button_save_e13 = tk.Button(ExcelFunction13_window, text='添加', font=('Arial', 12))
button_save_e13.place(x=70, y=200, width=90)
button_save_e13.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(merge_word_files, 13,
process_file1_path.get().strip(),process_file1_start.get().strip(),entry_path_file1_end.get().strip(),entry_path_file2_path.get().strip(),
entry_path_file2_start.get().strip(),entry_path_file2_end.get().strip(),entry_path_output_path.get().strip()
                                                    )))


"""----------Excel-merge_pdf_files------------"""
#  pdf1_path, pdf1_start, pdf1_end, pdf2_path, pdf2_start, pdf2_end, output_path

ExcelFunction14_window = tk.Frame(info_frame1)
pdf1_path_label = tk.Label(ExcelFunction14_window, text='pdf1文件路径', font=('Arial', 12))
pdf1_path_label.place(x=10, y=10)
pdf1_path_path= ttk.Entry(ExcelFunction14_window, font=('Arial', 12))
pdf1_path_path.place(x=105, y=10, width=115)

pdf1_start_label = tk.Label(ExcelFunction14_window, text='pdf1的起始页', font=('Arial', 12))
pdf1_start_label.place(x=10, y=35)
pdf1_start_start = ttk.Entry(ExcelFunction14_window, font=('Arial', 12))
pdf1_start_start.place(x=105, y=35, width=115)

pdf1_end = tk.Label(ExcelFunction14_window, text='pdf1的总结页', font=('Arial', 12))
pdf1_end.place(x=10, y=60)
entry_path_pdf1_end= ttk.Entry(ExcelFunction14_window, font=('Arial', 12))
entry_path_pdf1_end.place(x=105, y=60, width=115)

pdf2_path = tk.Label(ExcelFunction14_window, text='pdf2文件路径', font=('Arial', 12))
pdf2_path.place(x=10, y=85)
entry_path_pdf2_path= ttk.Entry(ExcelFunction14_window, font=('Arial', 12))
entry_path_pdf2_path.place(x=105, y=85, width=115)

pdf2_start = tk.Label(ExcelFunction14_window, text='pdf2的起始页', font=('Arial', 12))
pdf2_start.place(x=10, y=110)
entry_path_pdf2_start = ttk.Entry(ExcelFunction14_window, font=('Arial', 12))
entry_path_pdf2_start.place(x=105, y=110, width=115)

pdf2_end = tk.Label(ExcelFunction14_window, text='pdf2的总结页', font=('Arial', 12))
pdf2_end.place(x=10, y=135)
entry_path_pdf2_end = ttk.Entry(ExcelFunction14_window, font=('Arial', 12))
entry_path_pdf2_end.place(x=105, y=135, width=115)

output_path = tk.Label(ExcelFunction14_window, text='合并后文件路径', font=('Arial', 12))
output_path.place(x=10, y=160)
entry_path_output_path= ttk.Entry(ExcelFunction14_window, font=('Arial', 12))
entry_path_output_path.place(x=105, y=160, width=115)

button_save_e13 = tk.Button(ExcelFunction14_window, text='添加', font=('Arial', 12))
button_save_e13.place(x=70, y=200, width=90)
button_save_e13.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(merge_pdf_files, 14,
pdf1_path_path.get().strip(),pdf1_start_start.get().strip(),entry_path_pdf1_end.get().strip(),
entry_path_pdf2_path.get().strip(),entry_path_pdf2_start.get().strip(),entry_path_pdf2_end.get().strip(),entry_path_output_path.get().strip()
                                                    )))


"""----------Excel-EM_data_in_excel(最大期望法填充表格)------------"""

ExcelFunction15_window = tk.Frame(info_frame1)
fill_missing_data_in_excel_filepath_label = tk.Label(ExcelFunction15_window, text='excel文件路径', font=('Arial', 12))
fill_missing_data_in_excel_filepath_label.place(x=10, y=10)
fill_missing_data_in_excel_filepath = ttk.Entry(ExcelFunction15_window, font=('Arial', 12))
fill_missing_data_in_excel_filepath.place(x=95, y=10, width=115)

fill_missing_data_in_excel_sheetname_label = tk.Label(ExcelFunction15_window, text='表格名称', font=('Arial', 12))
fill_missing_data_in_excel_sheetname_label.place(x=10, y=35)
fill_missing_data_in_excel_sheetname = ttk.Entry(ExcelFunction15_window, font=('Arial', 12))
fill_missing_data_in_excel_sheetname.place(x=95, y=35, width=115)

left_top_coord = tk.Label(ExcelFunction15_window, text='左上角角坐标:', font=('Arial', 12))
left_top_coord.place(x=10, y=60)
entry_path_left_top_coord = ttk.Entry(ExcelFunction15_window, font=('Arial', 12))
entry_path_left_top_coord.place(x=95, y=60, width=115)

right_bottom_coord = tk.Label(ExcelFunction15_window, text='右下角坐标:', font=('Arial', 12))
right_bottom_coord.place(x=10, y=85)
entry_path_right_bottom_coord = ttk.Entry(ExcelFunction15_window, font=('Arial', 12))
entry_path_right_bottom_coord.place(x=95, y=85, width=115)

button_save_e14 = tk.Button(ExcelFunction15_window, text='添加', font=('Arial', 12))
button_save_e14.place(x=70, y=130, width=90)
button_save_e14.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(EM_data_in_excel, 15,
                                             fill_missing_data_in_excel_filepath.get().strip(),
                                     fill_missing_data_in_excel_sheetname.get().strip(),entry_path_left_top_coord.get().strip(),entry_path_right_bottom_coord.get().strip())))
"""-------------------------------------------"""

"""----------Excel-fill_blank_cells(中位数填充表格)------------"""

ExcelFunction16_window = tk.Frame(info_frame1)
fill_missing_data_in_excel_filepath_label = tk.Label(ExcelFunction16_window, text='excel文件路径', font=('Arial', 12))
fill_missing_data_in_excel_filepath_label.place(x=10, y=10)
fill_missing_data_in_excel_filepath = ttk.Entry(ExcelFunction16_window, font=('Arial', 12))
fill_missing_data_in_excel_filepath.place(x=95, y=10, width=115)

fill_missing_data_in_excel_sheetname_label = tk.Label(ExcelFunction16_window, text='表格名称', font=('Arial', 12))
fill_missing_data_in_excel_sheetname_label.place(x=10, y=35)
fill_missing_data_in_excel_sheetname = ttk.Entry(ExcelFunction16_window, font=('Arial', 12))
fill_missing_data_in_excel_sheetname.place(x=95, y=35, width=115)

left_top_coord = tk.Label(ExcelFunction16_window, text='左上角角坐标:', font=('Arial', 12))
left_top_coord.place(x=10, y=60)
entry_path_left_top_coord = ttk.Entry(ExcelFunction16_window, font=('Arial', 12))
entry_path_left_top_coord.place(x=95, y=60, width=115)

right_bottom_coord = tk.Label(ExcelFunction16_window, text='右下角坐标:', font=('Arial', 12))
right_bottom_coord.place(x=10, y=85)
entry_path_right_bottom_coord = ttk.Entry(ExcelFunction16_window, font=('Arial', 12))
entry_path_right_bottom_coord.place(x=95, y=85, width=115)

button_save_e14 = tk.Button(ExcelFunction16_window, text='添加', font=('Arial', 12))
button_save_e14.place(x=70, y=130, width=90)
button_save_e14.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(fill_blank_cells, 16,
                                             fill_missing_data_in_excel_filepath.get().strip(),
                                     fill_missing_data_in_excel_sheetname.get().strip(),entry_path_left_top_coord.get().strip(),entry_path_right_bottom_coord.get().strip())))
"""-------------------------------------------"""

"""----------Excel-fill_blank_cells1(平均数填充表格)------------"""

ExcelFunction17_window = tk.Frame(info_frame1)
fill_missing_data_in_excel_filepath_label = tk.Label(ExcelFunction17_window, text='excel文件路径', font=('Arial', 12))
fill_missing_data_in_excel_filepath_label.place(x=10, y=10)
fill_missing_data_in_excel_filepath = ttk.Entry(ExcelFunction17_window, font=('Arial', 12))
fill_missing_data_in_excel_filepath.place(x=95, y=10, width=115)

fill_missing_data_in_excel_sheetname_label = tk.Label(ExcelFunction17_window, text='表格名称', font=('Arial', 12))
fill_missing_data_in_excel_sheetname_label.place(x=10, y=35)
fill_missing_data_in_excel_sheetname = ttk.Entry(ExcelFunction17_window, font=('Arial', 12))
fill_missing_data_in_excel_sheetname.place(x=95, y=35, width=115)

left_top_coord = tk.Label(ExcelFunction17_window, text='左上角角坐标:', font=('Arial', 12))
left_top_coord.place(x=10, y=60)
entry_path_left_top_coord = ttk.Entry(ExcelFunction17_window, font=('Arial', 12))
entry_path_left_top_coord.place(x=95, y=60, width=115)

right_bottom_coord = tk.Label(ExcelFunction17_window, text='右下角坐标:', font=('Arial', 12))
right_bottom_coord.place(x=10, y=85)
entry_path_right_bottom_coord = ttk.Entry(ExcelFunction17_window, font=('Arial', 12))
entry_path_right_bottom_coord.place(x=95, y=85, width=115)

button_save_e14 = tk.Button(ExcelFunction17_window, text='添加', font=('Arial', 12))
button_save_e14.place(x=70, y=130, width=90)
button_save_e14.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(fill_blank_cells1, 17,
                                             fill_missing_data_in_excel_filepath.get().strip(),
                                     fill_missing_data_in_excel_sheetname.get().strip(),entry_path_left_top_coord.get().strip(),entry_path_right_bottom_coord.get().strip())))
"""-------------------------------------------"""

"""----------Excel-fill_blank_cells2(众数填充表格)------------"""

ExcelFunction18_window = tk.Frame(info_frame1)
fill_missing_data_in_excel_filepath_label = tk.Label(ExcelFunction18_window, text='excel文件路径', font=('Arial', 12))
fill_missing_data_in_excel_filepath_label.place(x=10, y=10)
fill_missing_data_in_excel_filepath = ttk.Entry(ExcelFunction18_window, font=('Arial', 12))
fill_missing_data_in_excel_filepath.place(x=95, y=10, width=115)

fill_missing_data_in_excel_sheetname_label = tk.Label(ExcelFunction18_window, text='表格名称', font=('Arial', 12))
fill_missing_data_in_excel_sheetname_label.place(x=10, y=35)
fill_missing_data_in_excel_sheetname = ttk.Entry(ExcelFunction18_window, font=('Arial', 12))
fill_missing_data_in_excel_sheetname.place(x=95, y=35, width=115)

left_top_coord = tk.Label(ExcelFunction18_window, text='左上角角坐标:', font=('Arial', 12))
left_top_coord.place(x=10, y=60)
entry_path_left_top_coord = ttk.Entry(ExcelFunction18_window, font=('Arial', 12))
entry_path_left_top_coord.place(x=95, y=60, width=115)

right_bottom_coord = tk.Label(ExcelFunction18_window, text='右下角坐标:', font=('Arial', 12))
right_bottom_coord.place(x=10, y=85)
entry_path_right_bottom_coord = ttk.Entry(ExcelFunction18_window, font=('Arial', 12))
entry_path_right_bottom_coord.place(x=95, y=85, width=115)

button_save_e14 = tk.Button(ExcelFunction18_window, text='添加', font=('Arial', 12))
button_save_e14.place(x=70, y=130, width=90)
button_save_e14.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(OperateCommand(fill_blank_cells2, 18,
                                             fill_missing_data_in_excel_filepath.get().strip(),
                                     fill_missing_data_in_excel_sheetname.get().strip(),entry_path_left_top_coord.get().strip(),entry_path_right_bottom_coord.get().strip())))
"""-------------------------------------------"""
"""----------Excel-if2------------"""
ExcelFunction19_window = tk.Frame(info_frame1)
if_excel_filepath_label = tk.Label(ExcelFunction19_window, text='excel文件路径', font=('Arial', 12))
if_excel_filepath_label.place(x=10, y=10)
if_excel_filepath = ttk.Entry(ExcelFunction19_window, font=('Arial', 12))
if_excel_filepath.place(x=95, y=10, width=115)

if_excel_sheetname_label = tk.Label(ExcelFunction19_window, text='表格名称', font=('Arial', 12))
if_excel_sheetname_label.place(x=10, y=35)
if_excel_sheetname = ttk.Entry(ExcelFunction19_window, font=('Arial', 12))
if_excel_sheetname.place(x=95, y=35, width=115)

search_value = tk.Label(ExcelFunction19_window, text='检索值:', font=('Arial', 12))
search_value.place(x=10, y=60)
entry_path_search_value = ttk.Entry(ExcelFunction19_window, font=('Arial', 12))
entry_path_search_value.place(x=95, y=60, width=115)


type_value = tk.Label(ExcelFunction19_window, text='检索值类型:', font=('Arial', 12))
type_value.place(x=10, y=85)
type_value = ttk.Combobox(ExcelFunction19_window)
type_value .place(x=95, y=85, width=115)
type_value['values'] = ('int', 'float', 'string')
button_save_e10 = tk.Button(ExcelFunction19_window, text='添加', font=('Arial', 12))
button_save_e10.place(x=70, y=140, width=90)
button_save_e10.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(LogicalCommand(greater_than,19,if_excel_filepath.get().strip(),if_excel_sheetname.get().strip(),
                                                    entry_path_search_value.get().strip(),type_value.get().strip())))
"""-------------------------------------------"""
"""----------Excel-if3------------"""
ExcelFunction20_window = tk.Frame(info_frame1)
if_excel_filepath_label = tk.Label(ExcelFunction20_window, text='excel文件路径', font=('Arial', 12))
if_excel_filepath_label.place(x=10, y=10)
if_excel_filepath = ttk.Entry(ExcelFunction20_window, font=('Arial', 12))
if_excel_filepath.place(x=95, y=10, width=115)

if_excel_sheetname_label = tk.Label(ExcelFunction20_window, text='表格名称', font=('Arial', 12))
if_excel_sheetname_label.place(x=10, y=35)
if_excel_sheetname = ttk.Entry(ExcelFunction20_window, font=('Arial', 12))
if_excel_sheetname.place(x=95, y=35, width=115)

search_value = tk.Label(ExcelFunction20_window, text='检索值:', font=('Arial', 12))
search_value.place(x=10, y=60)
entry_path_search_value = ttk.Entry(ExcelFunction20_window, font=('Arial', 12))
entry_path_search_value.place(x=95, y=60, width=115)


type_value = tk.Label(ExcelFunction20_window, text='检索值类型:', font=('Arial', 12))
type_value.place(x=10, y=85)
type_value = ttk.Combobox(ExcelFunction20_window)
type_value .place(x=95, y=85, width=115)
type_value['values'] = ('int', 'float', 'string')
button_save_e10 = tk.Button(ExcelFunction20_window, text='添加', font=('Arial', 12))
button_save_e10.place(x=70, y=140, width=90)
button_save_e10.config(
    command=lambda: flows[active_flow_id].commandcomponent.bind(LogicalCommand(less_than,20,if_excel_filepath.get().strip(),if_excel_sheetname.get().strip(),
                                                    entry_path_search_value.get().strip(),type_value.get().strip())))


"""-------------------------------------------"""
def run_command_click():
    flows[active_flow_id].text_pad_1.delete('1.0', tk.END)
    flows[active_flow_id].commandcomponent.click()  # 执行命令列表中的命令
    flows[active_flow_id].commandcomponent.should_stop = False  # 重置是否出现报错的值


def clear_command_click():
    flows[active_flow_id].text_pad.delete('1.0', tk.END)  # 清空命令流程窗口信息
    flows[active_flow_id].text_pad_1.delete('1.0', tk.END)  # 清空控制台信息
    flows[active_flow_id].commandcomponent.clear()


def back_step():
    flows[active_flow_id].text_pad.edit_undo()
    flows[active_flow_id].commandcomponent.remove_last()
    # 窗口页面的回退
    if len(flows[active_flow_id].commandcomponent.command_list) > 0:

        last_command = flows[active_flow_id].commandcomponent.command_list[-1]
        if isinstance(last_command, EndCommand):
            select_window(100)

        elif isinstance(last_command, LogicalCommand):

            if len(last_command.children_command_list) == 0:
                select_window(last_command.get_id())

            else:
                select_window(last_command.children_command_list[-1].get_id())

        elif isinstance(last_command, EndCommand):
            select_window(100)

        else:
            select_window(last_command.get_id())

    elif len(flows[active_flow_id].commandcomponent.command_list) == 0:
        select_window(100)



def next_step():
    # 显示内容的恢复
    flows[active_flow_id].text_pad.edit_redo()
    print(flows[active_flow_id].logical_status)
    if flows[active_flow_id].commandcomponent.command_list:
        if isinstance(flows[active_flow_id].commandcomponent.command_list[-1],LogicalCommand):
            flows[active_flow_id].logical_status = True

    if flows[active_flow_id].commandcomponent.record_list:
        if isinstance(flows[active_flow_id].commandcomponent.record_list[-1], LogicalCommand):
            command = flows[active_flow_id].commandcomponent.record_list.pop()
            flows[active_flow_id].commandcomponent.command_list.append(command)
            select_window(command.get_id())
            flows[active_flow_id].logical_status = True
            return
    if flows[active_flow_id].logical_status:
        if flows[active_flow_id].commandcomponent.command_list[-1].record_children_list:
            child_command = flows[active_flow_id].commandcomponent.command_list[-1].record_children_list.pop()
            flows[active_flow_id].commandcomponent.command_list[-1].children_command_list.append(child_command)
            select_window(flows[active_flow_id].commandcomponent.command_list[-1].children_command_list[-1].get_id())
        else:
            flows[active_flow_id].logical_status = False

    if not flows[active_flow_id].logical_status:
        command = flows[active_flow_id].commandcomponent.record_list.pop()
        flows[active_flow_id].commandcomponent.command_list.append(command)
        select_window(flows[active_flow_id].commandcomponent.command_list[-1].get_id())


def save_flow():
    file_path = filedialog.asksaveasfilename(defaultextension=".flow")
    if file_path:
        flows[active_flow_id].commandcomponent.save_flow(file_path)


def load_flow():
    file_path = filedialog.askopenfilename(filetypes=[("Flow files", "*.flow")])
    if file_path:
        flows[active_flow_id].commandcomponent.load_flow(file_path)

def run_command_in_flow(flow_id):
    flow = flows[flow_id]
    flow.commandcomponent.click()

def run_commands_in_flows():
    for flow_id in flows:
        command_thread = threading.Thread(target=run_command_in_flow, args=(flow_id,))
        command_thread.start()

info_frame2 = tk.LabelFrame(right_name, text='命令流程管理')
info_frame2.place(x=10, y=255, width=320, height=90)

button_back = tk.Button(info_frame2, text='撤消刚添加的命令', font=('Arial', 12))
button_back.grid(column=0, row=0, sticky=tk.EW, pady=5)
button_back.config(command=back_step)

button_next = tk.Button(info_frame2, text='恢复刚撤消的命令', font=('Arial', 12))
button_next.grid(column=1, row=0, sticky=tk.EW, pady=5)
button_next.config(command=next_step)

button_run_1 = tk.Button(info_frame2, text='运行命令流程', font=('Arial', 12))
button_run_1.grid(column=0, row=1, sticky=tk.EW, pady=5)
button_run_1.config(command=run_command_click)

button_clear_1 = tk.Button(info_frame2, text='清除命令流程', font=('Arial', 12))
button_clear_1.grid(column=1, row=1, sticky=tk.EW, pady=5)
button_clear_1.config(command=clear_command_click)

command_flow_frame = tk.LabelFrame(right_name, text='命令流程集管理')
command_flow_frame.place(x=10, y=345, width=320, height=140)

save_button = tk.Button(command_flow_frame, text="保存当前命令流程", command=save_flow)
save_button.grid(column=0, row=0, pady=5)

load_button = tk.Button(command_flow_frame, text="导入命令流程", command=load_flow)
load_button.grid(column=1, row=1, pady=5)

button_create_flow = tk.Button(command_flow_frame, text='创建新的命令流程', font=('Arial', 12))
button_create_flow.grid(column=0, row=1, pady=5)
button_create_flow.config(command=create_flow)

switch_flow_button = tk.Button(command_flow_frame, text="切换命令流程")
switch_flow_button.grid(column=1, row=0, pady=5)
switch_flow_button.config(command=switch_flow)

run_all_flow = tk.Button(command_flow_frame, text="运行所有命令流程")
run_all_flow.grid(column=0, row=2, pady=5)
run_all_flow.config(command=run_commands_in_flows)

run_all_flow = tk.Button(command_flow_frame, text="删除当前命令流程")
run_all_flow.grid(column=1, row=2,  pady=5)
run_all_flow.config(command=run_commands_in_flows)
root.mainloop()
